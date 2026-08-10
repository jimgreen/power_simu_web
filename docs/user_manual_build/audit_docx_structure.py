from __future__ import annotations

import argparse
import hashlib
import io
import re
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree
from PIL import Image


ROOT = Path(__file__).resolve().parent
DEFAULT_DOCX = ROOT / "manual_formula_final6.docx"

NS = {
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def text_of(node: etree._Element) -> str:
    return "".join(node.xpath(".//m:t/text() | .//w:t/text()", namespaces=NS)).strip()


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Audit the generated user-manual DOCX.")
    parser.add_argument(
        "docx",
        nargs="?",
        type=Path,
        default=DEFAULT_DOCX,
        help="DOCX candidate to audit",
    )
    parser.add_argument(
        "--final",
        type=Path,
        help="Optional final delivery copy; its SHA-256 must match the candidate",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    docx = args.docx.resolve()
    final_docx = args.final.resolve() if args.final else None

    assert docx.is_file(), docx
    if final_docx is not None:
        assert final_docx.is_file(), final_docx
        assert sha256(docx) == sha256(final_docx), "final DOCX differs from rendered candidate"

    with zipfile.ZipFile(docx) as archive:
        assert archive.testzip() is None, "ZIP CRC failure"
        names = archive.namelist()
        xml_names = [name for name in names if name.endswith(".xml")]
        xml_bytes = {name: archive.read(name) for name in xml_names}

        document_root = etree.fromstring(xml_bytes["word/document.xml"])
        math_nodes = document_root.xpath("//m:oMath", namespaces=NS)
        math_paragraphs = document_root.xpath("//m:oMathPara", namespaces=NS)
        assert math_nodes, "no native Word equations found"
        assert all(text_of(node) for node in math_nodes), "empty native equation found"

        empty_script_parts: list[str] = []
        for tag in ("sub", "sup"):
            for node in document_root.xpath(f"//m:{tag}", namespaces=NS):
                if not text_of(node):
                    empty_script_parts.append(tag)
        assert not empty_script_parts, f"empty math script nodes: {empty_script_parts}"

        empty_operands: list[str] = []
        for tag in ("e", "num", "den"):
            for node in document_root.xpath(f"//m:{tag}", namespaces=NS):
                if not text_of(node):
                    empty_operands.append(tag)
        assert not empty_operands, f"empty equation operands: {empty_operands[:20]}"

        full_xml = b"\n".join(xml_bytes.values()).decode("utf-8", errors="replace")
        forbidden = {
            "replacement character": "\ufffd",
            "visible square placeholder": "\u25a1",
            "OpenAI citation token": "\ue200cite\ue202",
            "internal citation marker": "turn0search",
            "temporary lorem ipsum": "Lorem ipsum",
        }
        hits = {label: token for label, token in forbidden.items() if token in full_xml}
        assert not hits, f"forbidden tokens found: {hits}"
        assert not re.search(r"\?{4,}", full_xml), "possible garbled text (four or more question marks)"

        media_names = [name for name in names if name.startswith("word/media/")]
        png_names = [name for name in media_names if name.lower().endswith(".png")]
        assert png_names, "no embedded PNG images found"
        image_sizes: list[tuple[int, int]] = []
        for name in png_names:
            payload = archive.read(name)
            with Image.open(io.BytesIO(payload)) as image:
                image.verify()
            with Image.open(io.BytesIO(payload)) as image:
                assert image.width > 0 and image.height > 0, name
                image_sizes.append((image.width, image.height))

    document = Document(docx)
    assert document.sections, "document has no sections"
    assert document.paragraphs, "document has no paragraphs"
    assert document.inline_shapes, "document has no inline pictures"

    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    required_headings = [
        "简介",
        "整体流程与并发架构",
        "模拟台功能介绍",
        "模拟台潮流模拟技术",
        "模拟台对外 WEB 接口",
        "学员台功能介绍",
        "学员台新能源优先控制技术",
        "总结与展望",
    ]
    missing = [heading for heading in required_headings if heading not in body_text]
    assert not missing, f"missing required chapters: {missing}"

    print("DOCX_PATH", docx)
    if final_docx is not None:
        print("FINAL_DOCX_PATH", final_docx)
    print("DOCX_SHA256", sha256(docx))
    print("ZIP_ENTRIES", len(names))
    print("XML_PARTS", len(xml_names))
    print("OMATH_NODES", len(math_nodes))
    print("OMATH_PARAGRAPHS", len(math_paragraphs))
    print("PARAGRAPHS", len(document.paragraphs))
    print("TABLES", len(document.tables))
    print("SECTIONS", len(document.sections))
    print("INLINE_SHAPES", len(document.inline_shapes))
    print("MEDIA_FILES", len(media_names))
    print("VALID_PNGS", len(png_names))
    print("PNG_DIMENSION_RANGE", min(image_sizes), max(image_sizes))
    print("STRUCTURE_AUDIT", "PASS")


if __name__ == "__main__":
    main()
