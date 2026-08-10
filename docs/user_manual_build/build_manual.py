from __future__ import annotations

import json
import math
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from lxml import etree
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\codex\power_simu_web")
BUILD = ROOT / "docs" / "user_manual_build"
SCREENSHOTS = BUILD / "screenshots"
GENERATED = BUILD / "generated"
OUTPUT = Path(
    os.environ.get(
        "USER_MANUAL_OUTPUT",
        str(ROOT / "docs" / "极地微电网模拟台与学员台用户使用手册.docx"),
    )
)

FONT_PATHS = [
    Path(r"C:\Windows\Fonts\msyh.ttc"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),
    Path(r"C:\Windows\Fonts\simsun.ttc"),
]
FONT_PATH = next((path for path in FONT_PATHS if path.exists()), FONT_PATHS[-1])

BLUE = "2E74B5"
NAVY = "17365D"
TEAL = "238B8E"
GREEN = "4F8A5B"
ORANGE = "D97734"
GOLD = "B58B2A"
RED = "B64B4B"
INK = "24313D"
MUTED = "647382"
LIGHT = "F4F7FA"
LIGHT_BLUE = "E8F1F8"
LIGHT_GREEN = "EAF4EC"
LIGHT_ORANGE = "FFF1E6"
GRID = "CCD7E0"


@dataclass
class PageSpec:
    title: str
    lead: str
    image: str | None = None
    caption: str = ""
    bullets: list[str] = field(default_factory=list)
    table_headers: list[str] = field(default_factory=list)
    table_rows: list[list[str]] = field(default_factory=list)
    table_widths: list[float] = field(default_factory=list)
    code: str = ""
    callout: str = ""
    source: str = ""
    image_height: float = 3.55
    equations: list[tuple[str, str]] = field(default_factory=list)


@dataclass
class ChapterSpec:
    number: int
    title: str
    pages: list[PageSpec]


SUPPLEMENTAL_CHECKS: dict[str, list[list[str]]] = {
    "主页：运行总览与能量流": [
        ["时钟与模型", "模型 ID、运行模式、仿真时刻和倍率一致，刷新页面不会推进仿真时钟。"],
        ["功率口径", "DC/AC 变流器总功率与总目标均按 P_DC 汇总；DC→AC 显示正值。"],
        ["绿电指标", "复核 P_green=P_L,dc+P_L,ac-P_diesel，且总负荷为零时不计算占比。"],
        ["异常状态", "量测过期、设备无效或死岛时显示告警，不以旧值继续冒充当前状态。"],
    ],
    "模型管理：新建、复制、导入和切换": [
        ["文件完整性", "模型、控制、量测、状态、曲线和接线图文件成套存在，导入后可重新解析。"],
        ["版本一致性", "所有写操作携带当前 revision；旧版本提交被拒绝并提示刷新。"],
        ["生命周期", "切换或替换模型时退休旧服务实例，旧潮流结果和旧控制代次不得回写。"],
        ["验收方法", "复制后修改副本并运行单步，原模型定义和运行目录保持不变。"],
    ],
    "电网模型：设备树、参数表与稳定身份": [
        ["稳定身份", "设备由模型块、idx、端子和显式关联索引定位；name 仅用于显示。"],
        ["真实边界", "柴发、储能、风光和变流器的额定值、上下限、SOC 与步长字段均可追溯。"],
        ["控制模式", "AC/DC 侧控制类型与可写设值字段一致，不由 dev_type 或名称推断。"],
        ["失效行为", "端子缺失、索引重复、关联越界或参数非法时闭锁相关设备并记录原因。"],
    ],
    "设备详情与 DC/AC 控制模式": [
        ["控直流", "dc_control_type=P 时生成 p_dc_set，目标值直接使用统一 P_DC 口径。"],
        ["控交流", "dc_control_type=NONE 且 AC 控制有效时生成 p_ac_set，写入值为 -P_DC。"],
        ["双 NONE", "两侧控制类型均为 NONE 时按约定回退 p_dc_set，但不修改模型控制类型。"],
        ["方向核对", "DC→AC：p_dc>0、p_ac<0；AC→DC：p_dc<0、p_ac>0。"],
    ],
    "接线图：拓扑定位与设备状态": [
        ["结构与运行态", "结构端子来自模型定义，开关和投退状态只叠加运行图，不反写源定义。"],
        ["参考源", "AC 岛必须有合格相角参考，DC 岛必须有电压参考，否则该岛 fail closed。"],
        ["死岛识别", "断开路径后的无源分量不计入新能源统计，也不生成自动控制目标。"],
        ["显示边界", "SVG 只负责可视化，潮流和控制算法不读取图中文字判断设备类型。"],
    ],
    "曲线设置：风、光、温度与仿真尺度": [
        ["量纲检查", "风速、辐照度、温度及时间点单位明确，输入值处于模型允许范围。"],
        ["能力响应", "提高风速或辐照后，相应设备最大可发随能力曲线变化且不超过额定值。"],
        ["版本刷新", "保存曲线后 revision 更新，下一仿真步使用新曲线而当前已冻结步不被篡改。"],
        ["复现实验", "记录曲线版本、随机扰动参数和仿真起点，可重复得到同一场景输入。"],
    ],
    "负荷曲线：逐设备需求与典型曲线": [
        ["逐设备输入", "每台负荷独立取值或插值，不把系统总量平均回填到设备。"],
        ["交直流分区", "交流负荷进入 AC 节点方程，直流负荷进入 DC 节点方程。"],
        ["边界检查", "负荷目标满足设备功率范围，异常点被拒绝或按明确规则降级。"],
        ["平衡复核", "修改曲线后单步运行，检查变流器、储能或柴发响应与两侧功率平衡。"],
    ],
    "参数配置、故障与运行模式": [
        ["参数优先级", "真实上下限、SOC、步长和保护带先于经济目标进入活动边界。"],
        ["故障区分", "设备故障改变投退/开关，量测故障改变值与有效位，两类故障分别审计。"],
        ["模式叠加", "运行模式只覆盖显式指定设备和字段，禁止按名称关键字批量覆盖控制类型。"],
        ["并发写入", "多客户端修改通过 revision 和模型锁串行提交，冲突请求不静默覆盖。"],
    ],
    "人工修改：定义修改、重试与恢复": [
        ["提交前", "核对模型 ID、revision、设备稳定身份、字段白名单和取值范围。"],
        ["提交中", "后台在目标模型锁内验证并生成新定义快照，失败不产生部分写入。"],
        ["禁止覆盖", "人工修改不得隐式改变 AC/DC 控制模式或以运行态覆盖源模型定义。"],
        ["恢复验证", "重试或重置后重新初始化并执行单步，确认定义签名和量测点表同步。"],
    ],
    "控制指令：当前有效、历史与取消": [
        ["字段与方向", "遥调字段由显式控制模式选择，变流器显示值与实际端口写入值按符号换算。"],
        ["有效期", "自动指令按有效期或 generation 管理；人工保持指令必须显式取消。"],
        ["完整替换", "同 strategy_id 的新 generation 用完整快照替换旧代次，缺省设备目标不会残留。"],
        ["执行审计", "同时记录请求值、限幅后执行值、来源、时间、撤销原因和设备响应。"],
    ],
    "实时量测：遥测、遥信与测点趋势": [
        ["点表签名", "名称定义与数值帧使用同一 definition_signature，签名不一致立即丢弃。"],
        ["顺序与长度", "遥测/遥信值严格按名称数组顺序排列，数量不一致不做位置猜测。"],
        ["有效性", "missing、invalid、死岛和不适用测点保留状态，不用零值掩盖异常。"],
        ["趋势分段", "历史按 run_id 和模型生命周期分段，重新开始后不拼接上一运行段。"],
    ],
    "运行日志：输入边界、控制响应和求解摘要": [
        ["输入记录", "天气、负荷、设备状态、控制快照和动态能力边界可追溯到同一仿真步。"],
        ["求解记录", "保存迭代次数、残差、收敛状态、worker PID、耗时和功率差额。"],
        ["控制记录", "保存策略代次、字段、请求值、执行值、限幅、撤销和拒绝原因。"],
        ["故障判据", "invalid、failed、数据过期和 fail-closed 必须显式告警，不归类为正常最优。"],
    ],
    "交互链接与标准训练流程": [
        ["初始化", "从模拟台当前模型生成链接，学员台下载定义并完成版本握手。"],
        ["开环", "接收稳定后先预览策略，检查设备边界、功率方向、平衡松弛和告警。"],
        ["闭环", "启动后在两台指令页核对 strategy_id、generation、执行值与下一帧响应。"],
        ["结束", "先停止自动控制并确认撤销当前代次，再停止接收和模拟台时钟。"],
    ],
    "一次仿真步的完整计算流程": [
        ["输入冻结", "天气、负荷、拓扑、状态、控制和 SOC 在提交 worker 前形成不可变快照。"],
        ["求解门槛", "只有收敛且输出结构有效的结果才允许更新 SOC、量测、历史和日志。"],
        ["原子发布", "同一仿真步的时钟、设备、量测和摘要一次替换，客户端不读取半更新状态。"],
        ["失败处理", "超时或异常终止 worker，保留上一有效状态并发布明确失败摘要。"],
    ],
    "设备定义、关联参数与运行叠加": [
        ["主设备", "模型块给出电气端子、状态、控制方式和基础功率/电压参数。"],
        ["关联参数", "风光储参数通过稳定索引关联主设备，索引不存在或重复即判无效。"],
        ["叠加规则", "weather/stat/control 只覆盖允许字段，不能新造设备或改变结构身份。"],
        ["边界来源", "优化和潮流均从真实模型参数构造限值，不以显示名称或默认猜测补值。"],
    ],
    "拓扑、母线合并和运行岛": [
        ["母线合并", "零阻抗支路和闭合开关按规则合并，普通支路仍作为网络元件。"],
        ["运行岛", "每个 AC/DC 分量分别检查参考源、有效设备和跨侧变流器连接。"],
        ["死岛闭锁", "无参考、端子错误或变流器无效的分量不参与求解与自动策略。"],
        ["缓存更新", "拓扑签名变化时重建节点和稀疏矩阵映射，避免沿用旧索引关系。"],
    ],
    "交流支路方程": [
        ["参数", "核对 r、x、b、tap、shift 的标幺/实际量纲和基准值，r=x=0 不直接求倒数。"],
        ["方向", "两端电流和功率均按本端流出定义，支路损耗等于两端有功之和。"],
        ["状态", "断开支路不进入 Ybus；零阻抗元件通过母线合并或专用约束处理。"],
        ["数值验收", "无源支路应满足 P_ij+P_ji≥0，变压器分接比改变后端口量符合预期。"],
    ],
    "交流节点网络方程与控制节点": [
        ["注入符号", "发电向节点注入为正、负荷消耗为负，网络计算功率与指定注入同号比较。"],
        ["节点类型", "PQ 指定 P/Q，PV 指定 P/V，平衡节点指定 V/相角并承担剩余失配。"],
        ["残差构造", "每个未知量对应独立方程，参考相角不重复进入状态向量。"],
        ["越限处理", "无功或电压控制达到上下限时切换有效约束，并记录节点类型变化。"],
    ],
    "直流支路与节点网络方程": [
        ["支路电流", "I_ij=(V_i-V_j)/r_ij，r_ij 必须为有效正值；方向按 i→j。"],
        ["端口功率", "P_ij=V_i I_ij，且 P_ij+P_ji=r_ij I_ij²≥0。"],
        ["节点平衡", "指定注入与导纳网络功率在同一基准和符号口径下形成残差。"],
        ["参考电压", "每个有效 DC 岛必须有合格电压参考，否则该岛不进入可执行结果。"],
    ],
    "DC/AC 变流器方程与功率方向": [
        ["方向", "DC→AC 时 P_dc>0、P_ac<0；AC→DC 时 P_dc<0、P_ac>0。"],
        ["损耗", "P_loss=P_dc+P_ac≥0，并受 r1/r2、电压和视在功率共同影响。"],
        ["网络耦合", "同一设备同时进入 DC 与 AC 节点残差，不能作为两个独立电源重复统计。"],
        ["控制约束", "仅由 ac_control_type/dc_control_type 决定方程与设值字段，运行态不得覆盖模型定义。"],
    ],
    "混联全局状态、残差和网络方程": [
        ["状态向量", "仅收录有效岛的非参考角、受控电压和变流器内部控制变量。"],
        ["残差维度", "AC、DC、变流器和控制方程数量与未知量一致，重复/缺失方程立即失败。"],
        ["岛级隔离", "健康岛可独立求解；无效岛闭锁并从全局索引和统计中剔除。"],
        ["结果映射", "求解结果按稳定索引写回对应设备，禁止按名称或数组偶然顺序回填。"],
    ],
    "Newton-Raphson、稀疏雅可比与线性求解": [
        ["初值", "优先使用上一收敛状态；模型或拓扑变化后重建并检查电压、相角初值范围。"],
        ["雅可比", "数值/解析雅可比的行列顺序与残差、状态向量完全一致。"],
        ["收敛", "同时检查残差范数、迭代次数、有限数和设备结果边界，不只检查求解器返回码。"],
        ["失败", "奇异、发散、超时或 NaN 时不发布伪结果，记录诊断并触发 worker 恢复。"],
    ],
    "量测生成、跨进程隔离与日志校核": [
        ["真值与 SCADA", "real 保存潮流真值，scada 在真值上叠加噪声与量测故障。"],
        ["点表顺序", "量测输出与 meas 定义逐项对齐，缺失项保留状态而不移动后续位置。"],
        ["进程边界", "worker 只返回可序列化结果，不直接修改 WEB 进程中的模型对象。"],
        ["守恒检查", "日志对比发电、负荷、储能、变流器端口和支路损耗，解释剩余功率差额。"],
    ],
    "模型版本与点表一致性": [
        ["模型签名", "拓扑或定义变化导致 model_version 更新，客户端必须重新获取定义。"],
        ["点表签名", "量测/控制名称变化导致 definition_signature 更新，旧值帧立即失效。"],
        ["数组契约", "名称和值长度及顺序必须一致，客户端不得按已知名称局部猜补错位值。"],
        ["恢复流程", "丢弃不一致帧，重新读取交互链接、定义和全量快照后再恢复增量接收。"],
    ],
    "主页与接收状态": [
        ["数据源", "交互链接、teacher_api_base、模型 ID 和版本签名与当前训练任务一致。"],
        ["接收健康", "检查帧龄、同帧时长、量测可用率、run_id 和通讯失败次数。"],
        ["统计口径", "主页变流器功率/目标统一 P_DC，直流送交流显示正值。"],
        ["闭锁条件", "未初始化、未接收、数据过期或定义不一致时禁止自动控制提交。"],
    ],
    "模型管理与模型初始化": [
        ["链接输入", "使用模拟台当前模型生成的交互链接，核对 teacher_api_base 和 model_id。"],
        ["定义下载", "一次获取模型、测点和控制定义，并校验 model_version 与 definition_signature。"],
        ["身份规则", "初始化保留模型块、稳定索引和端子关系，不根据名称或 dev_type 猜测角色。"],
        ["切换清理", "重新初始化或切换模型前停止接收、撤销当前 generation 并退休旧状态。"],
    ],
    "电网模型与接线图": [
        ["定义来源", "学员台使用初始化下载的模型块、端子和稳定索引，不直接修改模拟台源文件。"],
        ["运行叠加", "接收快照只更新当前值、状态和有效位，不覆盖 AC/DC 控制模式。"],
        ["拓扑核对", "接线图显示与后台运行岛一致，死岛资源不出现在可控集合。"],
        ["不一致处理", "定义签名变化后停止接收和自动控制，重新初始化后方可恢复。"],
    ],
    "曲线显示与实时量测": [
        ["增量顺序", "差量帧必须基于已确认序号；缺帧或序号回退时改取全量快照。"],
        ["历史分段", "曲线按 run_id、模型和接收 epoch 分段，避免跨运行拼接。"],
        ["控制对比", "计划曲线与实时曲线使用同一固定快照时间基准和功率符号。"],
        ["无效量测", "缺失或过期测点显示状态并闭锁相关策略，不使用最后值无限延拓。"],
    ],
    "控制指令与人工修改": [
        ["人工遥调", "核对设备身份、控制模式、字段、方向和上下限后再提交。"],
        ["自动策略", "独立记录 strategy_id、generation 和完整快照，不混写人工来源。"],
        ["冲突处理", "按模拟台有效指令规则裁决并记录最终执行值，不由浏览器最后写入者静默覆盖。"],
        ["结果核验", "在模拟台指令页和下一接收帧同时确认登记、执行及设备响应。"],
    ],
    "参数配置、运行日志与通讯故障": [
        ["新鲜度参数", "刷新周期、请求超时、帧龄上限和同帧上限共同决定接收有效性。"],
        ["故障阈值", "连续通讯失败达到阈值后停止接收，禁止继续生成或提交新策略。"],
        ["撤销动作", "接收丢失或数据过期时撤销当前 renewable_priority generation。"],
        ["定位顺序", "依次检查模拟台健康、模型 ID、交互链接、版本签名、网络和日志错误。"],
    ],
    "新能源控制主页面": [
        ["前置条件", "模型初始化、实时接收、帧新鲜度、量测有效性和拓扑解析全部通过。"],
        ["开闭环", "开环只预览；闭环提交前再次检查 lease、generation 和快照仍为当前。"],
        ["指标", "分别展示交/直侧资源与系统总计，变流器系统值始终使用 P_DC。"],
        ["审计", "每轮保留目标、边界、弃电、柴发、失衡、告警、迭代和下发结果。"],
    ],
    "学员台标准训练流程与故障恢复": [
        ["启动顺序", "模拟台选模并启动后，学员台初始化、开始接收，再进入开环策略预览。"],
        ["闭环门槛", "帧龄、量测、拓扑、设备边界和策略日志全部正常后才切换闭环。"],
        ["故障恢复", "通讯异常时先停止自动控制，检查服务、链接、版本和网络，再重新初始化。"],
        ["结束确认", "停止控制后核对 generation 已撤销，停止接收后再结束模拟台仿真。"],
    ],
    "三阶段计算总览": [
        ["阶段 1", "先构造安全活动边界并检查精确功率平衡是否可行。"],
        ["阶段 2", "仅在精确平衡不可行时求不可避免的最小失衡 δ*。"],
        ["阶段 3", "固定 δ* 后再最小化弃电、柴发和次要调节代价。"],
        ["结果门槛", "任一阶段输入无效或求解失败时不生成可执行设备命令。"],
    ],
    "阶段 2：最小功率平衡松弛": [
        ["启用条件", "阶段 1 在活动安全边界内无法精确满足 AC/DC 平衡。"],
        ["主要输出", "最小失衡 δ*、对应可行目标、失衡所在侧和求解状态。"],
        ["DC 优先", "只在已证明精确平衡不可行且存在有效网侧变流器时用于字典序选择。"],
        ["异常处理", "边界无效或无可行回退时标记 failed，相关岛不下发控制。"],
    ],
    "阶段 3：新能源优先经济目标": [
        ["第一优先", "在固定 δ* 的可行面上最小化新能源弃电总和。"],
        ["第二优先", "在同等弃电水平下继续压低柴发总出力。"],
        ["次要目标", "调节平滑和并联 SOC 重分配只用于等价解选择，不突破安全边界。"],
        ["输出审计", "报告设备目标、弃电、柴发、δ、活动边界、迭代次数与严格/回退状态。"],
    ],
    "SOC、强制动作、降额、步长与保护带": [
        ["低 SOC", "禁止继续放电；进入强制区后目标必须向充电方向修正。"],
        ["高 SOC", "禁止继续充电；进入强制区后目标必须向放电方向修正。"],
        ["动态边界", "SOC 降额后的充放功率与额定值、普通步长取交集。"],
        ["保护优先", "构网储能/柴发保护带优先于普通步长，必要覆盖必须进入告警与审计。"],
    ],
    "DC/AC 变流器与并联合理分配": [
        ["端口与字段", "统一目标 P_DC；控直流写 p_dc_set=P_DC，控交流写 p_ac_set=-P_DC，双 NONE 写 p_dc_set。"],
        ["两机算例", "额定容量 100/50 kW、系统目标 90 kW 时先分配 60/30 kW；两机 P_DC 合计仍为 90 kW。"],
        ["饱和重分配", "若 1 号上限为 50 kW，则其取 50 kW，剩余 40 kW 在 2 号裕度允许时转移给 2 号。"],
        ["闭锁条件", "控制模式、端子、边界或拓扑无效时闭锁对应设备/并联组，不用名称猜测恢复。"],
    ],
}


def rgb(hex_value: str) -> RGBColor:
    value = hex_value.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def hex_rgb(hex_value: str) -> tuple[int, int, int]:
    value = hex_value.lstrip("#")
    return int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(FONT_PATH), size=size, index=0)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont, width: int) -> list[str]:
    lines: list[str] = []
    for paragraph in str(text).splitlines() or [""]:
        current = ""
        for char in paragraph:
            trial = current + char
            if draw.textbbox((0, 0), trial, font=fnt)[2] <= width:
                current = trial
            else:
                if current:
                    lines.append(current)
                current = char
        lines.append(current)
    return lines


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    fnt: ImageFont.FreeTypeFont,
    fill: str | tuple[int, int, int],
    width: int,
    line_gap: int = 8,
    anchor: str | None = None,
) -> int:
    x, y = xy
    lines = wrap_text(draw, text, fnt, width)
    bbox = draw.textbbox((0, 0), "国Ag", font=fnt)
    line_h = bbox[3] - bbox[1] + line_gap
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill, anchor=anchor)
        y += line_h
    return y


def add_title_band(image: Image.Image, title: str, subtitle: str = "") -> ImageDraw.ImageDraw:
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, image.width, 116), fill=hex_rgb(NAVY))
    draw.text((52, 28), title, font=font(34, True), fill="white")
    if subtitle:
        draw.text((54, 76), subtitle, font=font(19), fill=hex_rgb("D8E6F1"))
    return draw


def save_diagram(name: str, title: str, subtitle: str, columns: Sequence[tuple[str, str, str]], footer: str = "") -> str:
    GENERATED.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1600, 900), "white")
    draw = add_title_band(img, title, subtitle)
    colors = [BLUE, TEAL, ORANGE, GREEN, GOLD, RED]
    n = len(columns)
    gap = 34
    margin = 58
    box_w = (1600 - 2 * margin - gap * (n - 1)) // n
    top = 185
    bottom = 720
    for index, (heading, body, accent) in enumerate(columns):
        x0 = margin + index * (box_w + gap)
        x1 = x0 + box_w
        color = accent or colors[index % len(colors)]
        draw.rounded_rectangle((x0, top, x1, bottom), radius=18, fill=hex_rgb(LIGHT), outline=hex_rgb(color), width=4)
        draw.rectangle((x0, top, x1, top + 78), fill=hex_rgb(color))
        draw.text((x0 + 24, top + 20), heading, font=font(26, True), fill="white")
        draw_wrapped(draw, (x0 + 24, top + 110), body, font(22), hex_rgb(INK), box_w - 48, line_gap=13)
        if index < n - 1:
            ax = x1 + 6
            ay = (top + bottom) // 2
            draw.line((ax, ay, ax + gap - 12, ay), fill=hex_rgb(MUTED), width=6)
            draw.polygon([(ax + gap - 12, ay), (ax + gap - 27, ay - 12), (ax + gap - 27, ay + 12)], fill=hex_rgb(MUTED))
    if footer:
        draw.rounded_rectangle((58, 760, 1542, 848), radius=14, fill=hex_rgb(LIGHT_BLUE), outline=hex_rgb(GRID), width=2)
        draw_wrapped(draw, (82, 782), footer, font(20), hex_rgb(NAVY), 1430, line_gap=8)
    path = GENERATED / name
    img.save(path, quality=95)
    return name


def save_card_grid(name: str, title: str, cards: Sequence[tuple[str, str, str]], subtitle: str = "") -> str:
    GENERATED.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1600, 900), "white")
    draw = add_title_band(img, title, subtitle)
    cols = 2 if len(cards) <= 4 else 3
    rows = math.ceil(len(cards) / cols)
    margin_x, gap_x = 60, 34
    margin_y, gap_y = 170, 28
    box_w = (1600 - 2 * margin_x - gap_x * (cols - 1)) // cols
    box_h = (850 - margin_y - gap_y * (rows - 1)) // rows
    for index, (heading, body, accent) in enumerate(cards):
        row, col = divmod(index, cols)
        x0 = margin_x + col * (box_w + gap_x)
        y0 = margin_y + row * (box_h + gap_y)
        x1, y1 = x0 + box_w, y0 + box_h
        draw.rounded_rectangle((x0, y0, x1, y1), radius=16, fill=hex_rgb(LIGHT), outline=hex_rgb(accent), width=4)
        draw.text((x0 + 22, y0 + 18), heading, font=font(25, True), fill=hex_rgb(accent))
        draw_wrapped(draw, (x0 + 22, y0 + 64), body, font(20), hex_rgb(INK), box_w - 44, line_gap=10)
    path = GENERATED / name
    img.save(path, quality=95)
    return name


def save_soc_chart(name: str) -> str:
    img = Image.new("RGB", (1600, 900), "white")
    draw = add_title_band(img, "SOC 边界、强制动作与功率降额", "功率正值表示放电，负值表示充电")
    left, right, top, bottom = 130, 1480, 190, 720
    draw.line((left, bottom, right, bottom), fill=hex_rgb(INK), width=4)
    draw.line((left, top, left, bottom), fill=hex_rgb(INK), width=4)
    for value, label in [(0.1, "SOC下限"), (0.2, "低SOC强充区"), (0.8, "高SOC强放区"), (0.9, "SOC上限")]:
        x = left + int((right - left) * value)
        draw.line((x, top, x, bottom), fill=hex_rgb(GRID), width=3)
        draw.text((x - 52, bottom + 20), label, font=font(18), fill=hex_rgb(MUTED))
    points_charge = [(left, top + 95), (left + 135, top + 95), (left + 270, top + 240), (right, top + 240)]
    points_discharge = [(left, top + 390), (left + 1080, top + 390), (left + 1215, top + 245), (right, top + 245)]
    draw.line(points_charge, fill=hex_rgb(BLUE), width=8, joint="curve")
    draw.line(points_discharge, fill=hex_rgb(ORANGE), width=8, joint="curve")
    draw.text((180, 220), "允许充电功率（低SOC侧降额）", font=font(24, True), fill=hex_rgb(BLUE))
    draw.text((920, 610), "允许放电功率（高SOC侧降额）", font=font(24, True), fill=hex_rgb(ORANGE))
    draw.text((132, 762), "SOC", font=font(22, True), fill=hex_rgb(INK))
    draw.rounded_rectangle((130, 800, 1480, 860), radius=12, fill=hex_rgb(LIGHT_GREEN))
    draw.text((160, 816), "边界先于经济目标：越界方向闭锁；保护区触发强制充/放；正常区再执行新能源优先优化。", font=font(21), fill=hex_rgb(GREEN))
    path = GENERATED / name
    img.save(path, quality=95)
    return name


def save_equation_sheet(
    name: str,
    title: str,
    subtitle: str,
    equations: Sequence[tuple[str, str, str]],
) -> str:
    """Create a visual formula index; the DOCX carries native OMML equations."""
    colors = [BLUE, TEAL, ORANGE, GREEN, GOLD, RED]
    cards = [
        (label, note or "标准公式以 Word 原生专业公式排版展示。", colors[index % len(colors)])
        for index, (label, _latex, note) in enumerate(equations)
    ]
    return save_card_grid(name, title, cards, subtitle)


def save_cover_collage() -> str:
    sources = [SCREENSHOTS / "01_simulator_home.png", SCREENSHOTS / "22_trainee_renewable_control.png"]
    images = [Image.open(path).convert("RGB") for path in sources]
    canvas = Image.new("RGB", (1600, 950), hex_rgb(LIGHT))
    draw = ImageDraw.Draw(canvas)
    draw.rectangle((0, 0, 1600, 90), fill=hex_rgb(NAVY))
    draw.text((54, 24), "模拟台 + 学员台", font=font(36, True), fill="white")
    for index, img in enumerate(images):
        target_w = 720
        target_h = 760
        img.thumbnail((target_w, target_h), Image.Resampling.LANCZOS)
        x = 55 + index * 770
        y = 135
        panel = Image.new("RGB", (720, 760), "white")
        px = (720 - img.width) // 2
        py = (760 - img.height) // 2
        panel.paste(img, (px, py))
        canvas.paste(panel, (x, y))
        draw.rounded_rectangle((x, y, x + 720, y + 760), radius=16, outline=hex_rgb(BLUE if index == 0 else TEAL), width=5)
        draw.text((x + 26, 858), "模拟台：场景、潮流与量测" if index == 0 else "学员台：接收、策略与控制", font=font(23, True), fill=hex_rgb(NAVY))
    path = GENERATED / "cover_collage.png"
    canvas.save(path, quality=95)
    return path.name


def generate_diagrams() -> None:
    save_cover_collage()
    save_card_grid("chapter_map.png", "全书结构", [
        ("1 简介", "产品定位、用户角色、端口与快速开始", BLUE),
        ("2 整体流程", "多进程、多线程、多客户端与五条关键数据链路", TEAL),
        ("3 模拟台功能", "模型、曲线、接线图、指令、量测和日志", ORANGE),
        ("4 潮流技术", "设备定义、交直流方程、Newton 法与量测生成", GREEN),
        ("5 WEB 接口", "设备、实时边界、遥测遥信、历史、遥调遥控与示例", GOLD),
        ("6 学员台功能", "接收、显示、控制、诊断与训练流程", BLUE),
        ("7 新能源控制", "三阶段优化、约束、策略代次与撤销", RED),
        ("8 总结展望", "验收清单、扩展方向与运行建议", TEAL),
    ])
    save_diagram("architecture.png", "整体部署与并发架构", "两个 WEB 角色可作为独立进程运行；计算工作单元与 HTTP 客户端并发解耦", [
        ("模拟台 WEB 进程", "ThreadingHTTPServer\n模型注册表与时钟线程\n快照、命令、外部 API", BLUE),
        ("潮流工作进程", "ProcessPoolExecutor(spawn)\n默认 1 个 worker\n30 s 超时后终止并重建", ORANGE),
        ("学员台 WEB 进程", "ThreadingHTTPServer\n实时交换线程池\n新能源控制后台线程池", TEAL),
        ("多 WEB 客户端", "浏览器按模型读取共享快照\nGET 并发读取\nPOST 在模型锁内原子提交", GREEN),
    ], "当前新能源控制执行单元是学员台进程内后台线程池，不是独立 OS 子进程；其输入/输出契约可保持不变地进一步进程化。")
    save_card_grid("concurrency_matrix.png", "进程、线程和客户端职责", [
        ("进程隔离", "模拟台与学员台可分别启动；潮流求解默认放在 spawn 子进程，避免求解阻塞 HTTP。", BLUE),
        ("请求并发", "ThreadingHTTPServer 为并发客户端提供独立请求线程；同一模型共享内存快照。", TEAL),
        ("后台线程", "时钟推进、实时交换、控制周期分别由后台线程触发，耗时任务交线程池。", ORANGE),
        ("一致性边界", "模型锁、定义 revision、service_instance_id、receive_epoch 和 generation 共同线性化写入。", GREEN),
    ])
    save_diagram("sim_flow_worker.png", "模拟台后台到潮流工作进程", "一次计算采用完整输入快照，返回求解结果和运行状态书", [
        ("冻结输入", "克隆 model/stat/weather/control\n计算曲线点与有效指令\n形成 SimulationConfig", BLUE),
        ("跨进程提交", "ProcessPoolExecutor.submit\n序列化完整配置\n记录提交/往返时间", ORANGE),
        ("混联潮流求解", "拓扑准备\nAC/DC/变流器残差\n稀疏 Newton 迭代", TEAL),
        ("原子发布", "返回 Snapshot、SOC、量测\n模型锁内替换上一帧\n追加历史与日志", GREEN),
    ])
    save_diagram("sim_web_flow.png", "模拟台后台与模拟台 WEB 前端", "前端不直接接触求解器，只通过版本化 JSON 快照读写", [
        ("GET 读取", "/api/snapshot\n/api/measurements/delta\n/api/runtime-logs", BLUE),
        ("共享快照", "clock + devices + measurements\ndefinitions + settings\nmodel_version", TEAL),
        ("POST 写入", "/api/clock\n/api/curves\n/api/device-faults\n/api/settings", ORANGE),
        ("前端刷新", "多客户端独立轮询\n差量优先、全量兜底\n失败不修改后台状态", GREEN),
    ])
    save_diagram("teacher_trainee_flow.png", "模拟台后台到学员台后台", "交互链接先交换模型/定义版本，再周期拉取运行帧和回送控制命令", [
        ("交互链接", "/api/trainee-link\nmodel_id + model_version\n各接口路径", BLUE),
        ("定义同步", "模型初始化下载定义包\n签名校验\n稳定索引与测点顺序", ORANGE),
        ("实时接收", "snapshot 或 delta\n帧龄与同帧检查\n本地发布 receive_epoch", TEAL),
        ("控制回送", "/api/student/commands\nstrategy_id + generation\n模拟台完整代次替换", GREEN),
    ])
    save_diagram("trainee_control_flow.png", "学员台后台、WEB 前端与新能源控制单元", "读取、预览、自动控制和下发通过同一模型级状态协调", [
        ("学员台 WEB", "查看接收状态与实时快照\n修改控制参数\n启动/停止自动控制", BLUE),
        ("实时交换层", "4 线程刷新池\n帧版本、帧龄、receive_epoch\n提供控制租约", TEAL),
        ("控制工作单元", "后台调度线程 + 4 线程执行池\n收集输入、三阶段优化\n构造完整策略快照", ORANGE),
        ("模拟台命令入口", "generation 原子替换\n按设备字段校验\n过期/停止/切换时撤销", GREEN),
    ])
    save_diagram("generation_lifecycle.png", "strategy_id + generation 生命周期", "每轮是完整快照，不在上一轮指令上做增量叠加", [
        ("生成 N", "读取固定 receive_epoch\n计算完整设备目标\n携带 replace_generation=true", BLUE),
        ("提交 N", "提交前再次验证服务实例、模型修订、接收代次与帧身份", TEAL),
        ("替换 N-1", "模拟台取消同 strategy_id 的旧代次，再登记 N 的完整目标集合", ORANGE),
        ("撤销 N", "停止、接收丢失、数据过期、模型切换、控制器退休均发送代次撤销", RED),
    ])
    save_diagram("simulation_cycle.png", "模拟台单步仿真计算流程", "仿真时钟与墙钟解耦；每个时步发布一帧一致性快照", [
        ("1 取边界", "按仿真时刻读取天气、负荷曲线\n合并人工修改与有效指令", BLUE),
        ("2 限值与状态", "投退/开合/死岛\n风光可发、柴发上下限\n储能 SOC 与功率边界", ORANGE),
        ("3 潮流求解", "构建拓扑和稀疏网络方程\n混联 Newton 迭代\n校验收敛", TEAL),
        ("4 发布结果", "更新 SOC\n生成 real/scada\n快照、历史、日志原子替换", GREEN),
    ])
    save_card_grid("model_blocks.png", "设备定义与稳定关系", [
        ("节点", "ACNode、DCNode：idx、额定电压、初值、运行状态。", BLUE),
        ("支路", "ACBranch/DCBranch、Break、ZeroBranch：i_node/j_node 形成稳定端子关系。", TEAL),
        ("电源与负荷", "AC/DC Generator、Load 及风光储关联块：通过显式索引引用主设备。", ORANGE),
        ("变流器", "DCACConverter/ACDCConverter：两端节点、损耗、控制模式、两端设值。", GREEN),
        ("运行叠加", "stat/control/weather 只覆盖同块、同 idx/name 的允许字段，不重写源模型。", GOLD),
        ("角色原则", "不按设备名称猜测，不依赖 dev_type 角色标记；拓扑或边界无效时 fail closed。", RED),
    ])
    save_card_grid("topology.png", "拓扑准备与运行岛过滤", [
        ("结构图", "节点、支路、开关、变流器端子形成无向连接图；稳定 idx 用于数组映射。", BLUE),
        ("运行图", "run_stat、开关状态和节点状态叠加后，仅保留当前有效边。", TEAL),
        ("分量/母线", "零阻抗边合并为计算母线，连通分量形成 AC/DC 运行岛。", ORANGE),
        ("参考点", "每个有效 AC 岛选择参考相角/平衡源；DC 岛选择电压参考。", GREEN),
        ("混联过滤", "AC/DC 岛通过有效网侧变流器关联，死岛设备不进入可控资源集合。", GOLD),
        ("失效策略", "端子缺失、重复稳定身份、不可达或无参考源时闭锁相关岛，不扩大到无关岛。", RED),
    ])
    save_card_grid("ac_equations.png", "交流支路与节点方程", [
        ("支路导纳", "y_ij = 1 / (r_ij + j x_ij)，并计入线路充电 b/2、变比 tap 和移相 shift。", BLUE),
        ("节点注入", "S_i = P_i + jQ_i = V_i · conj(Σ_j Y_ij V_j)。", TEAL),
        ("功率平衡", "ΔP_i = P_spec,i - P_calc,i；ΔQ_i = Q_spec,i - Q_calc,i。", ORANGE),
        ("节点类型", "PQ 同时列 ΔP/ΔQ；PV 固定 P 与电压幅值；Slack 固定幅值和相角。", GREEN),
    ])
    save_card_grid("dc_equations.png", "直流支路与节点方程", [
        ("支路电流", "I_ij = g_ij (V_i - V_j)，其中 g_ij = 1/r_ij。", BLUE),
        ("端口功率", "P_ij = V_i I_ij；线路损耗 P_loss = r_ij I_ij²。", TEAL),
        ("节点平衡", "ΔP_i = P_spec,i - V_i Σ_j G_ij V_j。", ORANGE),
        ("参考节点", "电压控制节点固定 V；其余节点以有功平衡方程求电压。", GREEN),
    ])
    save_card_grid("converter_equations.png", "DC/AC 变流器耦合与符号", [
        ("统一显示", "界面和系统统计统一使用 P_DC 口径：DC→AC 为正，AC→DC 为负。", BLUE),
        ("物理端口", "DC→AC：p_dc>0、p_ac<0；AC→DC：p_dc<0、p_ac>0。", TEAL),
        ("功率耦合", "理想情形 p_dc + p_ac = 0；考虑效率后，送端功率覆盖受端功率和损耗。", ORANGE),
        ("控制字段", "dc_control_type=P→p_dc_set；dc_control_type=NONE→p_ac_set；两侧 NONE 回退 p_dc_set。", GREEN),
    ])
    save_diagram("newton.png", "混联 Newton-Raphson 求解", "将 AC、DC 与网侧变流器方程装配为一个稀疏非线性系统", [
        ("状态向量 x", "AC 非参考相角与 PQ 电压\nDC 非参考电压\n变流器内部/控制变量", BLUE),
        ("残差 F(x)", "AC ΔP/ΔQ\nDC ΔP\n变流器功率/电压控制残差", ORANGE),
        ("稀疏雅可比 J", "预缓存行列模式\n每轮刷新数值\n组合 AC/DC/变流器块", TEAL),
        ("更新与判据", "J Δx = -F\nx←x+Δx\n||F||∞<tol 或达到最大迭代", GREEN),
    ])
    save_diagram("measurement_pipeline.png", "真值、SCADA 与历史量测", "测点定义顺序是名称/值接口和差量接口的一致性基础", [
        ("潮流 Snapshot", "设备端口功率、电压、电流\n投退/开合/SOC\n天气与时钟", BLUE),
        ("real 真值", "按 meas 定义逐点取值\n缺失点标记 invalid\n保持定义顺序", TEAL),
        ("scada 量测", "叠加随机噪声\n应用量测故障\n生成有效位", ORANGE),
        ("WEB 发布", "全量、差量、历史缓存\ndefinition_signature\n多客户端一致读取", GREEN),
    ])
    save_diagram("external_api_map.png", "模拟台对外 WEB 接口地图", "外部程序先读取交互链接，再使用返回的 external_api 路径", [
        ("发现", "/api/trainee-link\n模型 ID、版本签名\n接口路径清单", BLUE),
        ("设备与边界", "/external/devices\n/realtime-inputs\n天气与交直流负荷", TEAL),
        ("量测与历史", "/telemetry/names + values\n/values/query\n/history/query", ORANGE),
        ("控制", "/controls/names\n/controls/execute\n有效期、来源、撤销", GREEN),
    ])
    save_diagram("realtime_inputs_flow.png", "外部实时环境与负荷输入", "接口只更新模拟台运行曲线；潮流在相应仿真时刻读取，不由请求主动推进时钟", [
        ("1 发现与读取契约", "trainee-link.external_api.realtime_inputs\nGET 返回天气字段、单位、负荷键\n返回曲线点数与采样间隔", BLUE),
        ("2 选择输入模式", "single_point：单点\npoints：逐点对象数组\nseries：字段序列数组", TEAL),
        ("3 原子校验与写入", "校验时标、点数、间隔、字段、设备与范围\n任一错误整帧拒绝\n写入 runtime/curves.json", ORANGE),
        ("4 对潮流生效", "不修改原始 curves.json\n不推进仿真时钟\n下一点或未来对应时刻由潮流读取", GREEN),
    ])
    save_card_grid("realtime_inputs_contract.png", "实时输入的时序与原子性契约", [
        ("时标权威", "批量输入必须给出 start_time/绝对分钟/绝对秒之一；起点必须位于当前曲线采样网格。", BLUE),
        ("间隔约束", "point_interval_seconds 或 minutes 二选一，且必须是曲线采样间隔的正整数倍。", TEAL),
        ("数组一致", "points 长度等于 point_count 且字段集合一致；series 每条数组长度等于 point_count。", ORANGE),
        ("标准负荷键", "使用 ACLoad:设备名 或 DCLoad:设备名；裸名称仅在模型内唯一时接受。", GREEN),
        ("原子失败", "负数风速/辐照/负荷、非法湿度/气压、未知设备等返回 400，曲线和 revision 均不改变。", RED),
        ("并发边界", "若潮流正在计算，请求等待本轮完成后一次写入；已完成结果不会因曲线修订被丢弃。", GOLD),
    ])
    save_diagram("version_handshake.png", "模型与点表一致性握手", "定义变化必须触发重新发现，运行值变化不会改变模型签名", [
        ("读取链接", "保存 model_version.signature\n保存 revision\n保存接口路径", BLUE),
        ("读取名称", "保存 definition_signature\n保存遥测/遥信名称顺序", TEAL),
        ("读取值帧", "核对两个 signature\n核对名称和值长度\n不一致则丢弃", ORANGE),
        ("恢复", "重新获取链接和名称\n替换本地缓存\n继续读取新帧", GREEN),
    ])
    save_diagram("renewable_inputs.png", "新能源优先控制输入与可控集合", "先证明设备身份、拓扑、量测和边界有效，再进入优化", [
        ("固定快照", "clock/run_id/frame\nmodel revision/receive_epoch\nSCADA 与定义", BLUE),
        ("显式资源", "风/光/储/柴/网侧变流器\n模型块 + 稳定索引\n控制模式与设值字段", TEAL),
        ("运行拓扑", "AC/DC 分量\nDC 传输组\n结构接入与当前可达性", ORANGE),
        ("动态边界", "风光最大可发\nSOC 降额/强制动作\n步长、保护带、功率上下限", GREEN),
    ])
    save_diagram("three_stage.png", "三阶段新能源优先优化", "三阶段按字典序处理：先可行与安全，再最小失衡，最后在同等失衡下优化经济目标", [
        ("阶段 1", "构造变量和安全边界\n检查精确 AC/DC 平衡可行性\n必要时启用安全步长覆盖", RED),
        ("阶段 2", "在边界与并联约束内\n最小化功率平衡松弛\n变流冲突时 DC 优先、AC 次之", ORANGE),
        ("阶段 3", "固定阶段 2 的最小松弛\n最小弃电、最小柴发\n抑制无谓调节并按 SOC 再分配", GREEN),
    ])
    save_card_grid("stage1.png", "阶段 1：可行域与安全边界", [
        ("输入", "固定控制快照、拓扑岛、当前功率、可用功率、SOC、控制模式、保护参数。", BLUE),
        ("变量", "每台风光/柴发/储能的目标功率，以及网侧变流器统一 P_DC 目标。", TEAL),
        ("约束", "设备硬上下限、SOC 方向限制、动态降额、调节步长、并联变流器比例和拓扑岛归属。", ORANGE),
        ("可行性目标", "线性规划仅检查 A·P=A·P_current 是否存在；目标系数为 0。", GREEN),
        ("安全覆盖", "若强制 SOC 修正造成不可行，仅允许新能源向硬下限额外下调，并记录 step_override。", RED),
        ("输出", "active_lower/upper、精确平衡可行解或进入阶段 2 的最小失衡初值。", GOLD),
    ])
    save_card_grid("stage2.png", "阶段 2：最小功率失衡", [
        ("目标函数", "min ||δ||²，其中 δ = A·P_current - A·P；有 AC/DC 变流器冲突时先 min |δ_DC|，再固定其最优值 min |δ_AC|。", BLUE),
        ("等式约束", "并联变流器按额定可分配容量保持同比例；已证明可行时可直接保持 A·P=A·P_current。", TEAL),
        ("边界约束", "沿用阶段 1 的 active_lower ≤ P ≤ active_upper，不允许经济目标突破安全边界。", ORANGE),
        ("输入/输出", "输入安全可行域与当前值；输出最小松弛 δ*、一个可行 P* 和收敛状态。", GREEN),
    ])
    save_card_grid("stage3.png", "阶段 3：新能源优先经济目标", [
        ("目标函数", "min w_rΣ(P_avail-P_r)+w_dΣP_d+w_qΣ(P_avail-P_r)²+w_aΣ(P-P_current)²。", BLUE),
        ("平衡固定", "保持阶段 2 得到的 A·P=A·P_current-δ*；经济目标不能恶化最小失衡。", TEAL),
        ("优先次序", "通过正权重使新能源弃电最小、柴发出力最低；平方项仅用于平滑和打破等价解。", ORANGE),
        ("SOC 再分配", "在总平衡与边界不变的前提下，充电优先高 SOC 裕度设备，放电优先低 SOC 压力较小设备。", GREEN),
        ("输出", "每台设备目标、弃电、柴发目标、岛级 δ、迭代次数、状态、步长覆盖设备。", GOLD),
        ("失败处理", "数值失败或边界/并联校验失败时不下发该岛目标，保留诊断并 fail closed。", RED),
    ])
    save_soc_chart("soc_derating.png")
    save_diagram("parallel_converters.png", "并联变流器分配与控制字段", "系统总功率和总目标始终以 P_DC 统计；设备命令字段由控制模式决定", [
        ("组内总目标", "按 DC 传输组计算可调总量\n先裁剪到组总上下限\n正值表示 DC→AC", BLUE),
        ("容量分配", "并联设备按 allocation_capacity 比例分担\n边界饱和后按剩余裕度再分配", TEAL),
        ("端口换算", "P_DC 目标换算成 p_ac/p_dc 物理符号\n考虑变流方向与效率", ORANGE),
        ("字段选择", "dc=P → p_dc_set\ndc=NONE → p_ac_set\n双 NONE → p_dc_set", GREEN),
    ])
    save_diagram("strategy_snapshot.png", "策略输出、完整代次替换与撤销", "控制周期只发布已通过数据质量、边界和拓扑校验的设备", [
        ("计划结果", "command_rows + metrics + warnings\n岛级优化状态\n统一 P_DC 统计", BLUE),
        ("命令快照", "strategy_id=renewable_priority\ngeneration 单调递增\n完整 set_values 列表", TEAL),
        ("模拟台登记", "替换同策略上一代\n校验字段/上下限/有效期\n进入下一潮流时步", ORANGE),
        ("撤销与审计", "生命周期事件撤销当前代次\n日志保留原因、设备和时间\n不恢复旧代次", RED),
    ])
    save_card_grid("outlook.png", "总结与展望", [
        ("可操作", "两套控制台功能、接口和故障恢复流程统一在一册中。", BLUE),
        ("可解释", "潮流与三阶段优化均给出变量、方程、约束、输入、输出和失败语义。", TEAL),
        ("可集成", "外部接口具备版本握手、顺序一致性、历史查询和控制有效期。", ORANGE),
        ("可演进", "后续可将新能源控制工作单元进一步独立进程化，并增加认证、审计和高可用。", GREEN),
    ])
    save_equation_sheet(
        "power_conventions_math.png",
        "功率方向与绿电指标",
        "系统显示口径、物理端口符号和统计公式",
        [
            ("系统变流器口径", r"$P_{\mathrm{conv,sys}}=\sum_{c\in\mathcal C}P_{\mathrm{dc},c}$", "系统总功率和总目标统一使用各变流器直流端口功率之和。"),
            ("DC→AC", r"$P_{\mathrm{dc}}>0,\qquad P_{\mathrm{ac}}<0$", "直流侧送出功率，交流侧吸收负注入；界面显示为正。"),
            ("AC→DC", r"$P_{\mathrm{dc}}<0,\qquad P_{\mathrm{ac}}>0$", "交流侧送出功率，直流侧为负；界面显示为负。"),
            ("绿电指标", r"$P_{\mathrm{green}}=P_{L,\mathrm{dc}}+P_{L,\mathrm{ac}}-P_{\mathrm{diesel}},\qquad R_{\mathrm{green}}=\dfrac{P_{\mathrm{green}}}{P_{L,\mathrm{dc}}+P_{L,\mathrm{ac}}}\times100\%$", "分母为零时占比不计算。"),
        ],
    )
    save_equation_sheet(
        "ac_equations_math.png",
        "交流支路与节点网络方程",
        "复数相量形式，适用于线路、变压器和节点功率平衡",
        [
            ("串联导纳", r"$\underline{y}_{ij}=\dfrac{1}{r_{ij}+\mathrm{j}x_{ij}}=g_{ij}+\mathrm{j}b_{ij}$", "线路充电电纳按两端各 b_c/2 装配。"),
            ("支路电流", r"$\underline{I}_{ij}=\left(\underline{y}_{ij}+\mathrm{j}\dfrac{b_{c,ij}}{2}\right)\dfrac{\underline{V}_{i}}{|t_{ij}|^{2}}-\underline{y}_{ij}\dfrac{\underline{V}_{j}}{t_{ij}^{*}}$", "t_ij=τ_ij exp(jφ_ij) 为复变比；普通线路取 t_ij=1。"),
            ("节点复功率", r"$\underline{S}_{i}=P_i+\mathrm{j}Q_i=\underline{V}_{i}\left(\sum_{j=1}^{n}Y_{ij}\underline{V}_{j}\right)^{*}$", "Y 为运行拓扑对应的节点导纳矩阵。"),
            ("节点残差", r"$\Delta P_i=P_i^{\mathrm{sp}}-P_i(\boldsymbol{\theta},\boldsymbol{V}),\qquad \Delta Q_i=Q_i^{\mathrm{sp}}-Q_i(\boldsymbol{\theta},\boldsymbol{V})$", "PQ 节点使用两式，PV 节点固定 V 并使用有功平衡。"),
        ],
    )
    save_equation_sheet(
        "dc_equations_math.png",
        "直流支路与节点网络方程",
        "电导网络、端口功率和非参考节点有功平衡",
        [
            ("支路电流", r"$I_{ij}=g_{ij}(V_i-V_j),\qquad g_{ij}=\dfrac{1}{r_{ij}}$", "正方向定义为 i 节点流向 j 节点。"),
            ("端口功率", r"$P_{ij}=V_i I_{ij},\qquad P_{ji}=V_j I_{ji}$", "两端功率均按本端注入方向计算。"),
            ("线路损耗", r"$P_{\mathrm{loss},ij}=P_{ij}+P_{ji}=r_{ij}I_{ij}^{2}\ge0$", "损耗严格非负。"),
            ("节点残差", r"$\Delta P_i=P_i^{\mathrm{sp}}-V_i\sum_{j=1}^{n}G_{ij}V_j$", "非参考节点求 V_i；电压控制节点固定 V。"),
        ],
    )
    save_equation_sheet(
        "converter_equations_math.png",
        "DC/AC 变流器标准功率关系",
        "物理端口符号、r1/r2 损耗模型与系统 P_DC 口径",
        [
            ("端口方向", r"$P_{\mathrm{dc}}>0,\ P_{\mathrm{ac}}<0\ (\mathrm{DC}\to\mathrm{AC});\qquad P_{\mathrm{dc}}<0,\ P_{\mathrm{ac}}>0\ (\mathrm{AC}\to\mathrm{DC})$", "两端功率均按流入变流器为正，界面统一显示 P_DC。"),
            ("损耗方程", r"$V_{\mathrm{dc}}^2V_{\mathrm{ac}}^2(P_{\mathrm{dc}}+P_{\mathrm{ac}})-r_1P_{\mathrm{dc}}^2V_{\mathrm{ac}}^2-r_2(P_{\mathrm{ac}}^2+Q_{\mathrm{ac}}^2)V_{\mathrm{dc}}^2=0$", "与混联潮流内核的 r1/r2 端口损耗残差一致。"),
            ("非负损耗", r"$P_{\mathrm{loss}}=P_{\mathrm{dc}}+P_{\mathrm{ac}}=r_1(P_{\mathrm{dc}}/V_{\mathrm{dc}})^2+r_2(P_{\mathrm{ac}}^2+Q_{\mathrm{ac}}^2)/V_{\mathrm{ac}}^2\ge0$", "端口功率之和等于导通损耗。"),
            ("系统显示", r"$P_{\mathrm{conv,sys}}=\sum_{c\in\mathcal C}P_{\mathrm{dc},c}$", "无论命令写 p_dc_set 还是 p_ac_set，统计均先换算为 P_DC。"),
        ],
    )
    save_equation_sheet(
        "newton_math.png",
        "混联潮流 Newton-Raphson 标准形式",
        "AC、DC 与变流器方程构成统一稀疏非线性系统",
        [
            ("状态向量", r"$\boldsymbol{x}=\left[\boldsymbol{\theta}_{\mathrm{ac}}^{\mathsf T},\ \boldsymbol{V}_{\mathrm{ac}}^{\mathsf T},\ \boldsymbol{V}_{\mathrm{dc}}^{\mathsf T},\ \boldsymbol{x}_{\mathrm{conv}}^{\mathsf T}\right]^{\mathsf T}$", "仅包含非参考/非固定控制状态。"),
            ("残差向量", r"$\boldsymbol{F}(\boldsymbol{x})=\left[\Delta\boldsymbol{P}_{\mathrm{ac}}^{\mathsf T},\ \Delta\boldsymbol{Q}_{\mathrm{ac}}^{\mathsf T},\ \Delta\boldsymbol{P}_{\mathrm{dc}}^{\mathsf T},\ \boldsymbol{F}_{\mathrm{conv}}^{\mathsf T}\right]^{\mathsf T}$", "方程顺序与状态映射在 prepare 阶段固定。"),
            ("Newton 步", r"$\boldsymbol{J}(\boldsymbol{x}^{(k)})\Delta\boldsymbol{x}^{(k)}=\boldsymbol{F}(\boldsymbol{x}^{(k)}),\qquad \boldsymbol{x}^{(k+1)}=\boldsymbol{x}^{(k)}-\Delta\boldsymbol{x}^{(k)}$", "与程序求解 JΔx=F 后执行 x←x-Δx 的实现一致。"),
            ("收敛判据", r"$\left\|\boldsymbol{F}(\boldsymbol{x}^{(k)})\right\|_{\infty}\le\varepsilon_F$", "同时受最大迭代次数和线性求解成功约束。"),
        ],
    )
    save_equation_sheet(
        "capability_math.png",
        "风光储动态能力与 SOC 更新",
        "潮流求解前形成的物理可行边界",
        [
            ("风机可发", r"$P_w(v)=\begin{cases}0,&v<v_{\mathrm{ci}}\ \text{or}\ v\ge v_{\mathrm{co}}\\ P_r\left(\dfrac{v-v_{\mathrm{ci}}}{v_r-v_{\mathrm{ci}}}\right)^3,&v_{\mathrm{ci}}\le v<v_r\\ P_r,&v_r\le v<v_{\mathrm{co}}\end{cases}$", "切入、额定和切出风速来自设备参数。"),
            ("光伏可发", r"$P_{\mathrm{pv}}=P_r\dfrac{G}{G_{\mathrm{ref}}}\max\!\left\{0,1+\alpha_T(T-T_{\mathrm{ref}})\right\}$", "再裁剪到设备真实上下限。"),
            ("储能 SOC", r"$P_k^{\mathrm{int}}=\begin{cases}P_k/\eta_d,&P_k\ge0\\ \eta_cP_k,&P_k<0\end{cases},\qquad \mathrm{SOC}_{k+1}=\mathrm{SOC}_k-\dfrac{\Delta t}{E_r}P_k^{\mathrm{int}}$", "端口功率正值为放电、负值为充电；SOC 上下限通过可行功率边界保证。"),
        ],
    )
    save_equation_sheet(
        "measurement_math.png",
        "真值、SCADA 与有效位",
        "测点定义顺序不变，噪声和故障只作用于量测发布层",
        [
            ("测点映射", r"$z_i^{\mathrm{true}}=h_i(\boldsymbol{x},\boldsymbol{s}),\qquad i=1,\ldots,m$", "h_i 按 meas 定义从潮流状态和设备状态取值。"),
            ("噪声模型", r"$z_i^{\mathrm{scada}}=z_i^{\mathrm{true}}+\epsilon_i,\qquad \epsilon_i\sim\mathcal{N}(0,\sigma_i^2)$", "故障模型可进一步修改值或有效位。"),
            ("有效位", r"$q_i\in\{0,1\},\qquad q_i=0\Rightarrow z_i\ \text{不得参与控制}$", "missing、死岛或故障点均可置为无效。"),
        ],
    )
    save_equation_sheet(
        "stage1_math.png",
        "阶段 1：安全可行域与精确平衡",
        "零目标线性规划，仅证明安全边界内是否存在精确平衡解",
        [
            ("可行性模型", r"$\min_{\boldsymbol{P}\in\mathbb R^n}\ 0$", "这是标准线性可行性问题，本阶段不含弃电或柴发经济目标。"),
            ("功率平衡", r"$\boldsymbol{A}\boldsymbol{P}=\boldsymbol{A}\boldsymbol{P}^{0}$", "P0 为固定快照中的当前功率。"),
            ("并联约束", r"$\boldsymbol{C}_{\parallel}\boldsymbol{P}=\boldsymbol{0}$", "同 AC/DC 分量间并联变流器按可分配容量保持比例。"),
            ("安全边界", r"$\boldsymbol{\ell}^{\mathrm{act}}\le\boldsymbol{P}\le\boldsymbol{u}^{\mathrm{act}}$", "active 边界已合并硬上下限、SOC、降额、步长和保护修正。"),
        ],
    )
    save_equation_sheet(
        "stage2_math.png",
        "阶段 2：最小功率平衡松弛",
        "精确平衡不可行时，先最小化不可避免的 AC/DC 失衡",
        [
            ("松弛定义", r"$\boldsymbol{\delta}(\boldsymbol{P})=\boldsymbol{A}\boldsymbol{P}^{0}-\boldsymbol{A}\boldsymbol{P}$", "δ=0 表示保持本岛当前净平衡。"),
            ("一般目标", r"$\min_{\boldsymbol{P}}\ \left\|\boldsymbol{\delta}(\boldsymbol{P})\right\|_{2}^{2}$", "同时满足并联等式和 active 边界。"),
            ("DC 优先", r"$t_{\mathrm{dc}}^{*}=\min\ t_{\mathrm{dc}}\quad\mathrm{s.t.}\quad -t_{\mathrm{dc}}\le\delta_{\mathrm{dc}}(\boldsymbol{P})\le t_{\mathrm{dc}}$", "只在有网侧变流器且精确 AC/DC 平衡不可行时使用。"),
            ("AC 次优", r"$\min\ t_{\mathrm{ac}}\quad\mathrm{s.t.}\quad -t_{\mathrm{ac}}\le\delta_{\mathrm{ac}}(\boldsymbol{P})\le t_{\mathrm{ac}},\ |\delta_{\mathrm{dc}}(\boldsymbol{P})|\le t_{\mathrm{dc}}^{*}+\varepsilon$", "输出字典序最小松弛 δ*。"),
        ],
    )
    save_equation_sheet(
        "stage3_math.png",
        "阶段 3：新能源优先经济目标",
        "固定阶段 2 的最小松弛，在同一可行面上最小弃电和柴发",
        [
            ("目标函数", r"$\min_{\boldsymbol{P}}\ w_r\sum_{r\in\mathcal{R}}(P_r^{\mathrm{av}}-P_r)+w_d\sum_{d\in\mathcal{D}}P_d+w_q\sum_{r\in\mathcal{R}}(P_r^{\mathrm{av}}-P_r)^2+w_a\sum_{i\in\mathcal{S}\cup\mathcal{D}}(P_i-P_i^0)^2$", "前两项为主要优先级，平方项用于平滑和等价解选择。"),
            ("固定最小失衡", r"$\boldsymbol{A}\boldsymbol{P}=\boldsymbol{A}\boldsymbol{P}^{0}-\boldsymbol{\delta}^{*}$", "经济目标不得恶化阶段 2 的最小松弛。"),
            ("结构与边界", r"$\boldsymbol{C}_{\parallel}\boldsymbol{P}=\boldsymbol{0},\qquad \boldsymbol{\ell}^{\mathrm{act}}\le\boldsymbol{P}\le\boldsymbol{u}^{\mathrm{act}}$", "所有安全约束继续有效。"),
            ("主要指标", r"$P_{\mathrm{curt}}=\sum_{r\in\mathcal{R}}(P_r^{\mathrm{av}}-P_r),\qquad P_{\mathrm{diesel}}=\sum_{d\in\mathcal{D}}P_d$", "结果同时返回岛级弃电、柴发和松弛。"),
        ],
    )


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_inches: Sequence[float]) -> None:
    total_dxa = 9360
    widths = [int(round(value * 1440)) for value in widths_inches]
    if widths:
        widths[-1] += total_dxa - sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_inches[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, name: str = "Calibri", size: float | None = None, bold: bool | None = None, color: str | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = rgb(color)
    if italic is not None:
        run.italic = italic


def add_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_section(section, chapter_title: str) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.35)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run(chapter_title)
    set_run_font(hr, size=8.5, bold=True, color=MUTED)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fr = fp.add_run("极地微电网模拟台与学员台用户使用手册  |  第 ")
    set_run_font(fr, size=8, color=MUTED)
    add_field(fp, "PAGE")
    fr2 = fp.add_run(" 页  |  V1.0 · 2026-08-10")
    set_run_font(fr2, size=8, color=MUTED)


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Title", 31, NAVY, 0, 8),
        ("Subtitle", 13.5, MUTED, 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18
    if "Figure Caption" not in styles:
        cap = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Figure Caption"]
    cap.font.name = "Calibri"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    cap.font.size = Pt(8.5)
    cap.font.color.rgb = rgb(MUTED)
    cap.font.italic = True
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(5)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_kicker(doc: Document, text: str, page_break_before: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = page_break_before
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    set_run_font(r, size=8.5, bold=True, color=GOLD)


def add_lead(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=10.7, color=INK)


def add_picture(doc: Document, image_name: str, caption: str, max_height: float = 3.55) -> None:
    path = Path(image_name)
    if not path.is_absolute():
        candidate = SCREENSHOTS / image_name
        path = candidate if candidate.exists() else GENERATED / image_name
    if not path.exists():
        raise FileNotFoundError(path)
    with Image.open(path) as img:
        width_px, height_px = img.size
    width_in = 6.25
    height_in = width_in * height_px / width_px
    if height_in > max_height:
        height_in = max_height
        width_in = height_in * width_px / height_px
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in), height=Inches(height_in))
    if caption:
        cp = doc.add_paragraph(caption, style="Figure Caption")
        cp.paragraph_format.keep_with_next = True


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.keep_together = True
        r = p.add_run(str(item))
        set_run_font(r, size=10.25, color=INK)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=8.5, bold=True, color=NAVY)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2:
            for cell in cells:
                shade_cell(cell, "F8FAFC")
        for col_index, value in enumerate(row):
            p = cells[col_index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_run_font(r, size=8.15, color=INK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.02
    shade_paragraph(p, "F1F3F5")
    for index, line in enumerate(code.strip().splitlines()):
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=7.4, color=INK)
        if index < len(code.strip().splitlines()) - 1:
            r.add_break()


def add_callout(doc: Document, text: str, fill: str = LIGHT_BLUE, color: str = NAVY) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    shade_paragraph(p, fill)
    r = p.add_run(text)
    set_run_font(r, size=9.4, bold=True, color=color)


def add_source(doc: Document, text: str) -> None:
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("实现依据：" + text)
    set_run_font(r, size=7.4, color=MUTED, italic=True)


MATH_NS = "http://www.w3.org/1998/Math/MathML"
MML2OMML_XSL = Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL")
_MATH_TRANSFORM = None


def _m(tag: str, *children, text: str | None = None, **attrs):
    node = etree.Element(f"{{{MATH_NS}}}{tag}")
    for key, value in attrs.items():
        node.set(key.replace("_", "-"), str(value))
    if text is not None:
        node.text = str(text)
    for child in children:
        if child is None:
            continue
        node.append(child if hasattr(child, "tag") else _m("mtext", text=str(child)))
    return node


def mi(value: str, normal: bool = False):
    return _m("mi", text=value, mathvariant="normal" if normal else "italic")


def mn(value: str | int | float):
    return _m("mn", text=str(value))


def mo(value: str):
    return _m("mo", text=value)


def mt(value: str):
    return _m("mtext", text=value)


def row(*items):
    return _m("mrow", *items)


def sub(base, value):
    return _m("msub", base, value)


def sup(base, value):
    return _m("msup", base, value)


def subsup(base, low, high):
    return _m("msubsup", base, low, high)


def under(base, value):
    return _m("munder", base, value)


def frac(num, den):
    return _m("mfrac", num, den)


def fenced(content, left="(", right=")"):
    return _m("mfenced", content, open=left, close=right)


def summation(index_text: str, set_text: str):
    return under(mo("∑"), row(mi(index_text), mo("∈"), mi(set_text)))


def vector(value: str):
    return _m("mi", text=value, mathvariant="bold-italic")


def norm2(content):
    return sup(row(mo("‖"), content, mo("‖")), mn(2))


def norm_inf(content):
    return sub(row(mo("‖"), content, mo("‖")), mi("∞", True))


def display_math(*rows_):
    root = _m("math", display="block")
    if len(rows_) == 1:
        root.append(rows_[0])
    else:
        table = _m("mtable", columnalign="left")
        for item in rows_:
            table.append(_m("mtr", _m("mtd", item)))
        root.append(table)
    return etree.tostring(root, encoding="unicode")


def _math_transform():
    global _MATH_TRANSFORM
    if _MATH_TRANSFORM is None:
        if not MML2OMML_XSL.exists():
            raise FileNotFoundError(MML2OMML_XSL)
        _MATH_TRANSFORM = etree.XSLT(etree.parse(str(MML2OMML_XSL)))
    return _MATH_TRANSFORM


def add_equations(doc: Document, equations: Sequence[tuple[str, str]]) -> None:
    for label, mathml in equations:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(2)
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(label)
        set_run_font(lr, size=8.5, bold=True, color=BLUE)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        math_root = etree.fromstring(mathml.encode("utf-8"))
        omml = _math_transform()(math_root).getroot()
        p._p.append(omml)


def formula_registry() -> dict[str, tuple[str, str]]:
    P = lambda label: sub(mi("P"), mi(label, True))
    V = lambda label: sub(mi("V"), mi(label, True))
    registry: dict[str, tuple[str, str]] = {}

    registry["power_direction"] = (
        "式（1-1）变流器系统统计口径",
        display_math(row(subsup(mi("P"), mi("conv", True), mi("sys", True)), mo("≡"), P("dc"))),
    )
    registry["green_metric"] = (
        "式（1-2）绿电功率与绿电占比",
        display_math(
            row(P("green"), mo("="), P("L,dc"), mo("+"), P("L,ac"), mo("−"), P("diesel")),
            row(sub(mi("R"), mi("green", True)), mo("="), frac(P("green"), row(P("L,dc"), mo("+"), P("L,ac"))), mo("×"), mn(100), mo("%")),
        ),
    )
    registry["ac_admittance"] = (
        "式（4-1）交流支路串联导纳",
        display_math(row(sub(mi("y"), mi("ij", True)), mo("="), frac(mn(1), row(sub(mi("r"), mi("ij", True)), mo("+"), mi("j", True), sub(mi("x"), mi("ij", True)))), mo("="), sub(mi("g"), mi("ij", True)), mo("+"), mi("j", True), sub(mi("b"), mi("ij", True)))),
    )
    registry["ac_node_power"] = (
        "式（4-2）交流节点复功率方程",
        display_math(row(sub(mi("S"), mi("i", True)), mo("="), sub(mi("P"), mi("i", True)), mo("+"), mi("j", True), sub(mi("Q"), mi("i", True)), mo("="), sub(mi("V"), mi("i", True)), sup(fenced(row(mo("∑"), sub(mi("Y"), mi("ij", True)), sub(mi("V"), mi("j", True)))), mo("*")))),
    )
    registry["ac_residual"] = (
        "式（4-3）交流有功与无功残差",
        display_math(
            row(mi("ΔP", True), sub(mi(""), mi("i", True)), mo("="), subsup(mi("P"), mi("i", True), mi("sp", True)), mo("−"), sub(mi("P"), mi("i", True)), fenced(row(vector("θ"), mo(","), vector("V")))),
            row(mi("ΔQ", True), sub(mi(""), mi("i", True)), mo("="), subsup(mi("Q"), mi("i", True), mi("sp", True)), mo("−"), sub(mi("Q"), mi("i", True)), fenced(row(vector("θ"), mo(","), vector("V")))),
        ),
    )
    registry["dc_branch"] = (
        "式（4-4）直流支路电流、端口功率与损耗",
        display_math(
            row(sub(mi("I"), mi("ij", True)), mo("="), sub(mi("g"), mi("ij", True)), fenced(row(V("i"), mo("−"), V("j"))), mo(","), sub(mi("g"), mi("ij", True)), mo("="), frac(mn(1), sub(mi("r"), mi("ij", True)))),
            row(sub(mi("P"), mi("ij", True)), mo("="), V("i"), sub(mi("I"), mi("ij", True)), mo(","), subsup(mi("P"), mi("loss,ij", True), mi("")), mo("="), sub(mi("r"), mi("ij", True)), sup(sub(mi("I"), mi("ij", True)), mn(2)), mo("≥"), mn(0)),
        ),
    )
    registry["dc_residual"] = (
        "式（4-5）直流节点有功残差",
        display_math(row(mi("ΔP", True), sub(mi(""), mi("i", True)), mo("="), subsup(mi("P"), mi("i", True), mi("sp", True)), mo("−"), V("i"), row(mo("∑"), sub(mi("G"), mi("ij", True)), V("j")))),
    )
    registry["converter"] = (
        "式（4-6）DC/AC 变流器端口功率与效率",
        display_math(
            row(mt("DC→AC:"), P("ac"), mo("="), mo("−"), sub(mi("η"), mi("dca", True)), P("dc"), mo(","), P("dc"), mo(">"), mn(0)),
            row(mt("AC→DC:"), P("dc"), mo("="), mo("−"), sub(mi("η"), mi("acd", True)), P("ac"), mo(","), P("ac"), mo(">"), mn(0)),
            row(P("loss"), mo("="), P("dc"), mo("+"), P("ac"), mo("≥"), mn(0)),
        ),
    )
    registry["newton"] = (
        "式（4-7）混联潮流 Newton-Raphson 迭代",
        display_math(
            row(mi("J", True), fenced(sup(vector("x"), fenced(mi("k", True)))), mi("Δx", True), sup(fenced(mi("k", True)), mt("")), mo("="), mo("−"), mi("F", True), fenced(sup(vector("x"), fenced(mi("k", True))))),
            row(sup(vector("x"), fenced(row(mi("k", True), mo("+"), mn(1)))), mo("="), sup(vector("x"), fenced(mi("k", True))), mo("+"), sup(mi("Δx", True), fenced(mi("k", True)))),
            row(sub(norm2(mi("F", True),), mi("∞", True)), mo("≤"), sub(mi("ε"), mi("F", True))),
        ),
    )
    registry["pv"] = (
        "式（4-8）光伏最大可发功率",
        display_math(row(P("pv"), mo("="), P("r"), frac(mi("G"), sub(mi("G"), mi("ref", True))), mo("max"), fenced(row(mn(0), mo(","), mn(1), mo("+"), sub(mi("α"), mi("T", True)), fenced(row(mi("T"), mo("−"), sub(mi("T"), mi("ref", True)))))))),
    )
    wind_cases = _m(
        "mtable",
        _m("mtr", _m("mtd", mn(0)), _m("mtd", mt("v < v_ci 或 v ≥ v_co"))),
        _m(
            "mtr",
            _m("mtd", row(P("r"), sup(frac(row(mi("v"), mo("−"), sub(mi("v"), mi("ci", True))), row(sub(mi("v"), mi("r", True)), mo("−"), sub(mi("v"), mi("ci", True)))), mn(3)))),
            _m("mtd", mt("v_ci ≤ v < v_r")),
        ),
        _m("mtr", _m("mtd", P("r")), _m("mtd", mt("v_r ≤ v < v_co"))),
        columnalign="left left",
    )
    registry["wind_pv"] = (
        "式（4-8）风机分段功率曲线与光伏可发功率",
        display_math(
            row(P("w"), fenced(mi("v")), mo("="), _m("mfenced", wind_cases, open="{", close="")),
            row(P("pv"), mo("="), P("r"), frac(mi("G"), sub(mi("G"), mi("ref", True))), mo("max"), fenced(row(mn(0), mo(","), mn(1), mo("+"), sub(mi("α"), mi("T", True)), fenced(row(mi("T"), mo("−"), sub(mi("T"), mi("ref", True))))))),
        ),
    )
    registry["soc"] = (
        "式（4-9）储能 SOC 离散更新",
        display_math(row(sub(mi("SOC"), row(mi("k", True), mo("+"), mn(1))), mo("="), sub(mi("Π"), row(sub(mi("SOC"), mi("min", True)), mo(","), sub(mi("SOC"), mi("max", True)))), fenced(row(sub(mi("SOC"), mi("k", True)), mo("−"), frac(mi("Δt", True), sub(mi("E"), mi("r", True))), fenced(row(frac(sup(sub(mi("P"), mi("k", True)), mo("+")), sub(mi("η"), mi("d", True))), mo("+"), sub(mi("η"), mi("c", True)), sup(sub(mi("P"), mi("k", True)), mo("−")))))))),
    )
    registry["measurement"] = (
        "式（4-10）真值、SCADA 噪声与有效位",
        display_math(
            row(sub(mi("z"), mi("i,true", True)), mo("="), sub(mi("h"), mi("i", True)), fenced(row(vector("x"), mo(","), vector("s")))),
            row(sub(mi("z"), mi("i,scada", True)), mo("="), sub(mi("z"), mi("i,true", True)), mo("+"), sub(mi("ε"), mi("i", True)), mo(","), sub(mi("ε"), mi("i", True)), mo("∼"), mi("N", True), fenced(row(mn(0), mo(","), sup(sub(mi("σ"), mi("i", True)), mn(2))))),
            row(sub(mi("q"), mi("i", True)), mo("∈"), fenced(row(mn(0), mo(","), mn(1)), left="{", right="}"), mo(","), sub(mi("q"), mi("i", True)), mo("="), mn(0), mo("⇒"), mt("该点不得参与控制")),
        ),
    )
    registry["stage1"] = (
        "式（7-1）阶段 1 安全可行性模型",
        display_math(
            row(under(mi("min", True), vector("P")), mn(0)),
            row(mt("s.t."), vector("A"), vector("P"), mo("="), vector("A"), sup(vector("P"), mn(0))),
            row(vector("C"), sub(mi(""), mi("∥", True)), vector("P"), mo("="), vector("0")),
            row(sup(vector("ℓ"), mi("act", True)), mo("≤"), vector("P"), mo("≤"), sup(vector("u"), mi("act", True))),
        ),
    )
    registry["stage2"] = (
        "式（7-2）阶段 2 最小平衡松弛",
        display_math(
            row(vector("δ"), fenced(vector("P")), mo("="), vector("A"), sup(vector("P"), mn(0)), mo("−"), vector("A"), vector("P")),
            row(under(mi("min", True), vector("P")), sup(row(mo("‖"), vector("δ"), fenced(vector("P")), mo("‖")), mn(2))),
            row(mt("s.t."), vector("C"), sub(mi(""), mi("∥", True)), vector("P"), mo("="), vector("0"), mo(","), sup(vector("ℓ"), mi("act", True)), mo("≤"), vector("P"), mo("≤"), sup(vector("u"), mi("act", True))),
        ),
    )
    registry["stage3"] = (
        "式（7-3）阶段 3 新能源优先目标函数",
        display_math(
            row(
                under(mi("min", True), vector("P")),
                sub(mi("w"), mi("r", True)), row(mo("∑"), sub(row(fenced(row(sup(sub(mi("P"), mi("r", True)), mi("av", True)), mo("−"), sub(mi("P"), mi("r", True))))), row(mi("r", True), mo("∈"), mi("R", True)))),
                mo("+"), sub(mi("w"), mi("d", True)), row(mo("∑"), sub(sub(mi("P"), mi("d", True)), row(mi("d", True), mo("∈"), mi("D", True)))),
                mo("+"), sub(mi("w"), mi("q", True)), row(mo("∑"), sub(sup(fenced(row(sup(sub(mi("P"), mi("r", True)), mi("av", True)), mo("−"), sub(mi("P"), mi("r", True)))), mn(2)), row(mi("r", True), mo("∈"), mi("R", True)))),
                mo("+"), sub(mi("w"), mi("a", True)), row(mo("∑"), sub(sup(fenced(row(sub(mi("P"), mi("i", True)), mo("−"), sup(sub(mi("P"), mi("i", True)), mn(0)))), mn(2)), row(mi("i", True), mo("∈"), row(mi("S", True), mo("∪"), mi("D", True))))),
            ),
            row(mt("s.t."), vector("A"), vector("P"), mo("="), vector("A"), sup(vector("P"), mn(0)), mo("−"), sup(vector("δ"), mo("*"))),
            row(vector("C"), sub(mi(""), mi("∥", True)), vector("P"), mo("="), vector("0"), mo(","), sup(vector("ℓ"), mi("act", True)), mo("≤"), vector("P"), mo("≤"), sup(vector("u"), mi("act", True))),
        ),
    )
    return registry


def formula_registry_standard() -> dict[str, tuple[str, str]]:
    """Return publication-style equations without placeholder script nodes."""

    registry = formula_registry()
    P = lambda label: sub(mi("P"), mi(label, True))
    V = lambda label: sub(mi("V"), mi(label, True))
    # Use an ordinary scripted summation glyph. The Office MathML transform
    # otherwise creates an empty n-ary operand that LibreOffice renders as □.
    sum_n = lambda: subsup(mi("∑", True), row(mi("j"), mo("="), mn(1)), mi("n"))
    sum_set = lambda index, set_name: sub(mi("∑", True), row(mi(index), mo("∈"), mi(set_name)))
    max_set = lambda content: row(mi("max", True), fenced(content, left="{", right="}"))
    min_set = lambda content: row(mi("min", True), fenced(content, left="{", right="}"))

    p_conv_sys = lambda: sub(mi("P"), row(mi("conv", True), mo(","), mi("sys", True)))
    p_dc_c = lambda: sub(mi("P"), row(mi("dc", True), mo(","), mi("c")))
    p_ac_c = lambda: sub(mi("P"), row(mi("ac", True), mo(","), mi("c")))
    registry["power_direction"] = (
        "式（1-1）变流器系统统计口径与方向",
        display_math(
            row(p_conv_sys(), mo("="), sum_set("c", "C"), p_dc_c()),
            row(mt("DC→AC:"), p_dc_c(), mo(">"), mn(0), mo(","), p_ac_c(), mo("<"), mn(0)),
            row(mt("AC→DC:"), p_dc_c(), mo("<"), mn(0), mo(","), p_ac_c(), mo(">"), mn(0)),
        ),
    )
    registry["green_metric"] = (
        "式（1-2）绿电功率与绿电占比",
        display_math(
            row(P("green"), mo("="), P("L,dc"), mo("+"), P("L,ac"), mo("−"), P("diesel")),
            row(
                sub(mi("R"), mi("green", True)),
                mo("="),
                frac(P("green"), row(P("L,dc"), mo("+"), P("L,ac"))),
                mo("×"),
                mn(100),
                mo("%"),
                mo(","),
                P("L,dc"),
                mo("+"),
                P("L,ac"),
                mo(">"),
                mn(0),
            ),
        ),
    )
    registry["ac_residual"] = (
        "式（4-3）交流节点有功与无功残差",
        display_math(
            row(
                sub(mi("ΔP"), mi("i")),
                mo("="),
                sup(sub(mi("P"), mi("i")), mi("sp", True)),
                mo("−"),
                sub(mi("P"), mi("i")),
                fenced(row(vector("θ"), mo(","), vector("V"))),
            ),
            row(
                sub(mi("ΔQ"), mi("i")),
                mo("="),
                sup(sub(mi("Q"), mi("i")), mi("sp", True)),
                mo("−"),
                sub(mi("Q"), mi("i")),
                fenced(row(vector("θ"), mo(","), vector("V"))),
            ),
        ),
    )
    registry["dc_branch"] = (
        "式（4-4）直流支路电流、端口功率与损耗",
        display_math(
            row(
                sub(mi("I"), mi("ij")),
                mo("="),
                sub(mi("g"), mi("ij")),
                fenced(row(V("i"), mo("−"), V("j"))),
                mo(","),
                sub(mi("g"), mi("ij")),
                mo("="),
                frac(mn(1), sub(mi("r"), mi("ij"))),
            ),
            row(
                sub(mi("P"), mi("ij")),
                mo("="),
                V("i"),
                sub(mi("I"), mi("ij")),
                mo(","),
                sub(mi("P"), row(mi("loss", True), mo(","), mi("ij"))),
                mo("="),
                sub(mi("P"), mi("ij")),
                mo("+"),
                sub(mi("P"), mi("ji")),
                mo("="),
                sub(mi("r"), mi("ij")),
                sup(sub(mi("I"), mi("ij")), mn(2)),
                mo("≥"),
                mn(0),
            ),
        ),
    )
    registry["dc_residual"] = (
        "式（4-5）直流节点功率平衡残差",
        display_math(
            row(
                sup(sub(mi("F"), mi("i")), mi("dc", True)),
                mo("="),
                V("i"),
                sum_n(),
                sub(mi("G"), row(mi("i"), mi("j"))),
                V("j"),
                mo("+"),
                sup(sub(mi("I"), mi("i")), mi("sh", True)),
                V("i"),
                mo("−"),
                sup(sub(mi("P"), mi("i")), mi("sp", True)),
                mo("="),
                mn(0),
            )
        ),
    )
    q_ac = lambda: sub(mi("Q"), mi("ac", True))
    registry["converter"] = (
        "式（4-6）DC/AC 变流器端口功率与损耗方程",
        display_math(
            row(
                sup(V("dc"), mn(2)),
                sup(V("ac"), mn(2)),
                fenced(row(P("dc"), mo("+"), P("ac"))),
                mo("−"),
                sub(mi("r"), mn(1)),
                sup(P("dc"), mn(2)),
                sup(V("ac"), mn(2)),
                mo("−"),
                sub(mi("r"), mn(2)),
                fenced(row(sup(P("ac"), mn(2)), mo("+"), sup(q_ac(), mn(2)))),
                sup(V("dc"), mn(2)),
                mo("="),
                mn(0),
            ),
            row(
                P("loss"),
                mo("="),
                P("dc"),
                mo("+"),
                P("ac"),
                mo("="),
                sub(mi("r"), mn(1)),
                sup(frac(P("dc"), V("dc")), mn(2)),
                mo("+"),
                sub(mi("r"), mn(2)),
                frac(row(sup(P("ac"), mn(2)), mo("+"), sup(q_ac(), mn(2))), sup(V("ac"), mn(2))),
                mo("≥"),
                mn(0),
            ),
        ),
    )
    x_k = lambda: sup(vector("x"), fenced(mi("k")))
    dx_k = lambda: sup(vector("Δx"), fenced(mi("k")))
    f_x_k = lambda: row(vector("F"), fenced(x_k()))
    registry["newton"] = (
        "式（4-7）混联潮流 Newton-Raphson 迭代",
        display_math(
            row(vector("J"), fenced(x_k()), dx_k(), mo("="), f_x_k()),
            row(sup(vector("x"), fenced(row(mi("k"), mo("+"), mn(1)))), mo("="), x_k(), mo("−"), dx_k()),
            row(norm_inf(f_x_k()), mo("≤"), sub(mi("ε"), mi("F", True)), mo(","), vector("J"), mo("="), frac(row(mi("∂"), vector("F")), row(mi("∂"), vector("x")))),
        ),
    )
    wind_cases = _m(
        "mtable",
        _m("mtr", _m("mtd", mn(0)), _m("mtd", row(mi("v"), mo("<"), sub(mi("v"), mi("ci", True)), mo("∨"), mi("v"), mo("≥"), sub(mi("v"), mi("co", True))))),
        _m(
            "mtr",
            _m("mtd", row(P("r"), sup(frac(row(mi("v"), mo("−"), sub(mi("v"), mi("ci", True))), row(sub(mi("v"), mi("r", True)), mo("−"), sub(mi("v"), mi("ci", True)))), mn(3)))),
            _m("mtd", row(sub(mi("v"), mi("ci", True)), mo("≤"), mi("v"), mo("<"), sub(mi("v"), mi("r", True)))),
        ),
        _m("mtr", _m("mtd", P("r")), _m("mtd", row(sub(mi("v"), mi("r", True)), mo("≤"), mi("v"), mo("<"), sub(mi("v"), mi("co", True))))),
        columnalign="left left",
    )
    registry["wind_pv"] = (
        "式（4-8）风机与光伏最大可发功率",
        display_math(
            row(sup(P("w"), mi("av", True)), fenced(mi("v")), mo("="), _m("mfenced", wind_cases, open="{", close="")),
            row(
                sup(P("pv"), mi("av", True)),
                mo("="),
                P("r"),
                frac(mi("G"), sub(mi("G"), mi("ref", True))),
                max_set(row(mn(0), mo(","), mn(1), mo("+"), sub(mi("α"), mi("T", True)), fenced(row(mi("T"), mo("−"), sub(mi("T"), mi("ref", True)))))),
            ),
        ),
    )
    registry["soc"] = (
        "式（4-9）储能端口功率、内部功率与 SOC 更新",
        display_math(
            row(
                sup(sub(mi("P"), mi("k")), mi("cell", True)),
                mo("="),
                frac(sup(sub(mi("P"), mi("k")), mo("+")), sub(mi("η"), mi("d", True))),
                mo("+"),
                sub(mi("η"), mi("c", True)),
                sup(sub(mi("P"), mi("k")), mo("−")),
            ),
            row(
                sup(sub(mi("P"), mi("k")), mo("+")),
                mo("="),
                max_set(row(sub(mi("P"), mi("k")), mo(","), mn(0))),
                mo(","),
                sup(sub(mi("P"), mi("k")), mo("−")),
                mo("="),
                min_set(row(sub(mi("P"), mi("k")), mo(","), mn(0))),
            ),
            row(sub(mi("SOC"), row(mi("k"), mo("+"), mn(1))), mo("="), sub(mi("SOC"), mi("k")), mo("−"), frac(mi("Δt"), sub(mi("E"), mi("r", True))), sup(sub(mi("P"), mi("k")), mi("cell", True))),
            row(sub(mi("SOC"), mi("min", True)), mo("≤"), sub(mi("SOC"), mi("k")), mo("≤"), sub(mi("SOC"), mi("max", True))),
        ),
    )

    q_i = lambda: sub(mi("q"), mi("i"))
    i_ctrl = lambda: sub(mi("I"), mi("ctrl", True))
    registry["measurement"] = (
        "式（4-10）真值、SCADA 噪声与有效位",
        display_math(
            row(sup(sub(mi("z"), mi("i")), mi("true", True)), mo("="), sub(mi("h"), mi("i")), fenced(row(vector("x"), mo(","), vector("s")))),
            row(
                sup(sub(mi("z"), mi("i")), mi("scada", True)),
                mo("="),
                sup(sub(mi("z"), mi("i")), mi("true", True)),
                mo("+"),
                sub(mi("ε"), mi("i")),
                mo(","),
                sub(mi("ε"), mi("i")),
                mo("∼"),
                mi("N", True),
                fenced(row(mn(0), mo(","), sup(sub(mi("σ"), mi("i")), mn(2)))),
            ),
            row(
                q_i(),
                mo("∈"),
                fenced(row(mn(0), mo(","), mn(1)), left="{", right="}"),
                mo(","),
                i_ctrl(),
                mo("="),
                fenced(row(mi("i"), mo("|"), q_i(), mo("="), mn(1)), left="{", right="}"),
            ),
        ),
    )

    soc_i = lambda: sub(mi("SOC"), mi("i"))
    soc_min_i = lambda: sub(mi("SOC"), row(mi("min", True), mo(","), mi("i")))
    soc_max_i = lambda: sub(mi("SOC"), row(mi("max", True), mo(","), mi("i")))
    registry["safety_bounds"] = (
        "式（7-1）降额、保护带与活动安全边界",
        display_math(
            row(
                sup(P("ch,i"), mi("soc", True)),
                mo("="),
                sub(mi("γ"), mi("ch,i", True)),
                fenced(soc_i()),
                sup(P("ch,i"), mi("r", True)),
                mo(","),
                sup(P("dis,i"), mi("soc", True)),
                mo("="),
                sub(mi("γ"), mi("dis,i", True)),
                fenced(soc_i()),
                sup(P("dis,i"), mi("r", True)),
            ),
            row(
                sub(mi("g"), mi("d")),
                mo("="),
                min_set(row(sub(mi("ρ"), mi("d")), sup(P("d"), mi("max", True)), mo(","), frac(row(sup(P("d"), mi("max", True)), mo("−"), sup(P("d"), mi("min", True))), mn(2)))),
            ),
            row(
                sup(sub(mi("ℓ"), mi("i")), mi("act", True)),
                mo("="),
                max_set(row(sup(sub(mi("ℓ"), mi("i")), mi("safe", True)), mo(","), sup(P("i"), mn(0)), mo("−"), sup(sub(mi("ΔP"), mi("i")), mi("max", True)))),
                mo(","),
                sup(sub(mi("u"), mi("i")), mi("act", True)),
                mo("="),
                min_set(row(sup(sub(mi("u"), mi("i")), mi("safe", True)), mo(","), sup(P("i"), mn(0)), mo("+"), sup(sub(mi("ΔP"), mi("i")), mi("max", True)))),
            ),
            row(
                soc_i(), mo("<"), soc_min_i(), mo("−"), sub(mi("d"), mi("i")), mo("⇒"), P("i"), mo("<"), mn(0),
                mo(","),
                soc_i(), mo(">"), soc_max_i(), mo("+"), sub(mi("d"), mi("i")), mo("⇒"), P("i"), mo(">"), mn(0),
            ),
        ),
    )
    c_parallel = lambda: sub(vector("C"), mi("∥", True))
    registry["stage1"] = (
        "式（7-2）阶段 1 精确平衡可行性模型",
        display_math(
            row(under(mi("min", True), row(vector("P"), mo("∈"), sup(mi("ℝ", True), mi("n")))), mn(0)),
            row(mt("s.t."), vector("A"), vector("P"), mo("="), vector("A"), sup(vector("P"), mn(0))),
            row(c_parallel(), vector("P"), mo("="), vector("0")),
            row(sup(vector("ℓ"), mi("act", True)), mo("≤"), vector("P"), mo("≤"), sup(vector("u"), mi("act", True))),
        ),
    )
    delta_p = lambda: row(vector("δ"), fenced(vector("P")))
    registry["stage2"] = (
        "式（7-3）阶段 2 最小平衡松弛与 DC/AC 字典序",
        display_math(
            row(delta_p(), mo("="), vector("A"), sup(vector("P"), mn(0)), mo("−"), vector("A"), vector("P")),
            row(under(mi("min", True), vector("P")), norm2(delta_p()), mo(","), mt("s.t."), c_parallel(), vector("P"), mo("="), vector("0"), mo(","), sup(vector("ℓ"), mi("act", True)), mo("≤"), vector("P"), mo("≤"), sup(vector("u"), mi("act", True))),
            row(
                sup(sub(mi("t"), mi("dc", True)), mo("*")), mo("="), under(mi("min", True), row(vector("P"), mo(","), sub(mi("t"), mi("dc", True)))), sub(mi("t"), mi("dc", True)),
                mo(","), mo("−"), sub(mi("t"), mi("dc", True)), mo("≤"), sub(mi("δ"), mi("dc", True)), fenced(vector("P")), mo("≤"), sub(mi("t"), mi("dc", True)),
            ),
            row(
                under(mi("min", True), row(vector("P"), mo(","), sub(mi("t"), mi("ac", True)))), sub(mi("t"), mi("ac", True)),
                mo(","), mo("−"), sub(mi("t"), mi("ac", True)), mo("≤"), sub(mi("δ"), mi("ac", True)), fenced(vector("P")), mo("≤"), sub(mi("t"), mi("ac", True)),
                mo(","), row(mo("|"), sub(mi("δ"), mi("dc", True)), fenced(vector("P")), mo("|")), mo("≤"), sup(sub(mi("t"), mi("dc", True)), mo("*")), mo("+"), mi("τ"),
            ),
        ),
    )
    p_r = lambda: sub(mi("P"), mi("r"))
    p_r_av = lambda: sup(p_r(), mi("av", True))
    registry["stage3"] = (
        "式（7-4）阶段 3 新能源优先经济目标",
        display_math(
            row(under(mi("min", True), vector("P")), mi("J"), fenced(vector("P"))),
            row(
                mi("J"), fenced(vector("P")), mo("="),
                sub(mi("w"), mi("r")), sum_set("r", "R"), fenced(row(p_r_av(), mo("−"), p_r())),
                mo("+"), sub(mi("w"), mi("d")), sum_set("d", "D"), sub(mi("P"), mi("d")),
            ),
            row(
                mo("+"), sub(mi("w"), mi("q")), sum_set("r", "R"), sup(fenced(row(p_r_av(), mo("−"), p_r())), mn(2)),
                mo("+"), sub(mi("w"), mi("a")), sub(mi("∑", True), row(mi("i"), mo("∈"), row(mi("S"), mo("∪"), mi("D")))), sup(fenced(row(sub(mi("P"), mi("i")), mo("−"), sup(sub(mi("P"), mi("i")), mn(0)))), mn(2)),
            ),
            row(mt("s.t."), vector("A"), vector("P"), mo("="), vector("A"), sup(vector("P"), mn(0)), mo("−"), sup(vector("δ"), mo("*"))),
            row(c_parallel(), vector("P"), mo("="), vector("0"), mo(","), sup(vector("ℓ"), mi("act", True)), mo("≤"), vector("P"), mo("≤"), sup(vector("u"), mi("act", True))),
        ),
    )
    field_cases = _m(
        "mtable",
        _m("mtr", _m("mtd", mi("dc", True)), _m("mtd", row(sub(mi("d"), mi("c")), mo("="), mi("P", True), mo("∨"), fenced(row(sub(mi("a"), mi("c")), mo(","), sub(mi("d"), mi("c")))), mo("="), fenced(row(mi("NONE", True), mo(","), mi("NONE", True)))))),
        _m("mtr", _m("mtd", mi("ac", True)), _m("mtd", row(sub(mi("d"), mi("c")), mo("="), mi("NONE", True), mo(","), sub(mi("a"), mi("c")), mo("∈"), fenced(row(mi("PQ", True), mo(","), mi("PV", True), mo(","), mi("PH", True)), left="{", right="}")))),
        columnalign="left left",
    )
    command_cases = _m(
        "mtable",
        _m("mtr", _m("mtd", sub(mi("P"), mi("c"))), _m("mtd", row(sub(mi("φ"), mi("c")), mo("="), mi("dc", True)))),
        _m("mtr", _m("mtd", row(mo("−"), sub(mi("P"), mi("c")))), _m("mtd", row(sub(mi("φ"), mi("c")), mo("="), mi("ac", True)))),
        columnalign="left left",
    )
    registry["converter_dispatch"] = (
        "式（7-5）并联变流器分配、系统汇总与控制字段",
        display_math(
            row(frac(sub(mi("P"), mi("c")), sub(mi("K"), mi("c"))), mo("="), frac(sub(mi("P"), mi("c0", True)), sub(mi("K"), mi("c0", True))), mo("⇔"), sub(mi("K"), mi("c0", True)), sub(mi("P"), mi("c")), mo("−"), sub(mi("K"), mi("c")), sub(mi("P"), mi("c0", True)), mo("="), mn(0)),
            row(p_conv_sys(), mo("="), sum_set("c", "C"), p_dc_c()),
            row(sub(mi("φ"), mi("c")), mo("="), _m("mfenced", field_cases, open="{", close="")),
            row(sub(mi("u"), mi("c")), mo("="), _m("mfenced", command_cases, open="{", close="")),
        ),
    )
    return registry


def add_manual_page(doc: Document, chapter: ChapterSpec, page: PageSpec, first_in_section: bool) -> None:
    add_kicker(doc, f"第 {chapter.number} 章  {chapter.title}", page_break_before=not first_in_section)
    heading = doc.add_paragraph(page.title, style="Heading 1")
    heading.paragraph_format.space_before = Pt(2)
    add_lead(doc, page.lead)
    if page.image:
        add_picture(doc, page.image, page.caption, page.image_height)
    if page.table_headers:
        add_table(doc, page.table_headers, page.table_rows, page.table_widths)
    if page.code:
        add_code(doc, page.code)
    if page.equations:
        add_equations(doc, page.equations)
    if page.bullets:
        add_bullets(doc, page.bullets)
    supplemental_rows = SUPPLEMENTAL_CHECKS.get(page.title)
    if supplemental_rows:
        add_table(doc, ["核验项", "判定标准"], supplemental_rows, [1.35, 5.15])
    if page.callout:
        add_callout(doc, page.callout)
    add_source(doc, page.source)


def front_matter(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section, "极地微电网模拟台与学员台用户使用手册")

    add_kicker(doc, "极地微电网仿真与控制平台")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("模拟台与学员台\n用户使用手册")
    set_run_font(r, size=29, bold=True, color=NAVY)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(12)
    sr = sub.add_run("功能操作 · 并发架构 · 潮流模拟 · 外部接口 · 新能源优先控制")
    set_run_font(sr, size=12.8, color=MUTED)
    add_picture(doc, "cover_collage.png", "模拟台负责场景与潮流，学员台负责接收、研判和控制。", 5.0)
    add_callout(doc, "适用版本：power_simu_web 当前工作区实现  |  文档版本：V1.0  |  日期：2026-08-10", LIGHT_GREEN, GREEN)

    add_kicker(doc, "文档说明", page_break_before=True)
    doc.add_paragraph("阅读对象、范围与使用边界", style="Heading 1")
    add_lead(doc, "本手册将模拟台和学员台合并说明，既可用于日常操作，也可用于接口联调、算法评审和课堂训练。")
    add_picture(doc, "chapter_map.png", "八章结构覆盖从启动操作到算法和接口的完整链路。", 4.45)
    add_table(doc, ["读者", "重点章节", "建议用途"], [
        ["教员/运维", "1、2、3、6、8", "启动服务、建模、设置场景、监视运行和处理故障"],
        ["学员/调度员", "1、3、6、7", "接收数据、查看潮流、生成策略、下发并复盘"],
        ["算法工程师", "2、4、7", "核对方程、约束、数据生命周期和优化结果"],
        ["集成开发者", "2、5", "接入设备/遥测/遥信/历史/遥调遥控接口"],
    ], [1.15, 1.55, 3.8])

    add_kicker(doc, "目录与快速导航", page_break_before=True)
    doc.add_paragraph("八章导航", style="Heading 1")
    add_lead(doc, "全书按用户指定顺序编排；技术章节均从实际代码路径提取，并明确当前实现与可扩展部署的差异。")
    add_table(doc, ["章", "名称", "内容焦点"], [
        ["1", "简介", "系统定位、用户角色、端口、快速开始和功率口径"],
        ["2", "整体流程与并发架构", "多进程、多线程、多客户端及五条数据链路"],
        ["3", "模拟台功能介绍", "模型、接线、曲线、参数、指令、量测和日志"],
        ["4", "模拟台潮流模拟技术", "设备定义、支路/网络方程、Newton 法和计算流程"],
        ["5", "模拟台对外 WEB 接口", "版本握手、设备、量测、历史、控制和示例"],
        ["6", "学员台功能介绍", "模型初始化、实时接收、操作、诊断和新能源页面"],
        ["7", "学员台新能源优先控制技术", "三阶段目标函数、约束、输入输出和策略代次"],
        ["8", "总结与展望", "能力总结、验收清单、部署与安全演进"],
    ], [0.55, 2.25, 3.7])
    add_callout(doc, "阅读建议：首次使用按 1→2→3→6；算法评审按 4→7；外部系统联调直接进入第 5 章。", LIGHT_ORANGE, ORANGE)
    add_table(doc, ["常用任务", "直接阅读"], [
        ["启动并检查服务", "第 1 章快速启动；第 2 章运行单元；第 3/6 章主页状态"],
        ["建立训练链路", "第 3 章交互链接；第 6 章初始化、接收与标准训练流程"],
        ["分析潮流结果", "第 4 章设备、支路、网络、变流器、Newton 法与量测生成"],
        ["接入外部系统", "第 5 章接口发现、签名握手、实时/历史量测和控制示例"],
        ["评审新能源策略", "第 7 章三阶段目标、约束、SOC、保护带、步长和并联分配"],
        ["处理通讯或模型异常", "第 2 章 fail closed；第 6 章通讯故障；第 8 章验收清单"],
    ], [1.75, 4.75])
    add_table(doc, ["全书统一约定", "口径"], [
        ["设备身份", "模型块、稳定 idx、显式关联索引和端子关系；名称只作显示"],
        ["变流器功率", "系统总功率/总目标统一 P_DC；DC→AC 为正，AC→DC 为负"],
        ["失效策略", "拓扑、量测、边界、控制模式或身份无效时 fail closed"],
        ["自动控制", "strategy_id + generation；每轮完整快照替换并按生命周期撤销"],
    ], [1.75, 4.75])


def build_chapters() -> list[ChapterSpec]:
    F = formula_registry_standard()
    return [
        ChapterSpec(1, "简介", [
            PageSpec("系统定位与两台分工", "平台以同一微电网模型为基础，将场景生成、潮流求解和训练控制分离为两个用户角色。", "architecture.png", "模拟台、潮流工作单元、学员台与多客户端的总体关系。", [
                "模拟台是教员侧/环境侧：维护模型、天气与负荷曲线、故障、运行模式和仿真时钟，并发布真值与 SCADA。",
                "学员台是观察与控制侧：初始化本地模型、接收模拟台实时帧、查看量测和拓扑、生成并下发控制策略。",
                "两个 WEB 服务默认端口分别为 8710 和 8720，可作为独立进程运行；同一服务允许多个浏览器客户端并发访问。",
            ], source="simu/server.py::parse_args、make_http_server"),
            PageSpec("功率符号、统计口径与绿电指标", "手册统一采用当前界面口径，避免 AC/DC 与 DC/AC 名称差异造成方向误读。", "power_conventions_math.png", "功率方向与绿电指标定义。", [
                "系统变流器总功率/总目标统一使用 P_DC：DC→AC 为正，AC→DC 为负。",
                "物理端口满足：DC→AC 时 p_ac<0、p_dc>0；AC→DC 时 p_ac>0、p_dc<0。",
                "绿电功率 = 直流负荷功率 + 交流负荷功率 - 柴发功率；绿电占比 = 绿电功率 / 总负荷功率 × 100%。",
                "界面可将 ACDC/DCAC 名称统一显示为 DC/AC 变流器，但内部仍按模型块和控制模式选择实际设值字段。",
            ], image_height=2.0, equations=[F["power_direction"], F["green_metric"]], source="simu/device_roles.py、WEB 首页统计逻辑"),
            PageSpec("快速启动、端口与首次检查", "首次使用先确认两个后台、模型和潮流工作进程健康，再打开浏览器。", "01_simulator_home.png", "模拟台主页：当前模型、仿真模式、时钟、能量流和绿电指标。", [
                "模拟台默认访问 http://127.0.0.1:8710/；学员台默认访问 http://127.0.0.1:8720/。",
                "模拟台先选择模型并检查接线图、曲线和参数；学员台通过交互链接完成模型初始化后再启动接收。",
                "健康接口 /api/health 可检查角色、当前计算摘要以及潮流 worker PID、超时和重启次数。",
                "若服务绑定 127.0.0.1，只能本机访问；局域网访问应绑定受控网卡并配置防火墙白名单。",
            ], callout="不要把服务直接暴露到不可信网络：当前接口支持 CORS 和 gzip，但没有内置身份认证。", source="simu/server.py::main、/api/health"),
        ]),
        ChapterSpec(2, "整体流程与并发架构", [
            PageSpec("整体部署与运行单元", "并发能力分为进程隔离、后台线程、HTTP 请求线程和浏览器客户端四层。", "architecture.png", "默认部署结构及当前新能源控制工作单元的位置。", [
                "模拟台与学员台可以作为两个独立 OS 进程启动，各自维护模型注册表、运行目录和 WEB 静态资源。",
                "潮流求解默认由 ProcessPoolExecutor(spawn) 子进程执行，避免数值计算阻塞模拟台 HTTP 请求。",
                "新能源控制当前是学员台进程内的后台调度线程 + 4 线程执行池；手册按“控制工作单元”描述其数据契约。",
                "浏览器客户端不保存权威状态，只通过版本化快照读取；写操作在后台模型锁内串行提交。",
            ], source="simu/power_flow_worker.py、simu/server.py、simu/renewable_control.py"),
            PageSpec("多进程、多线程与多客户端矩阵", "同一进程可以服务多个模型和多个 WEB 客户端，但不同并发层承担的职责不同。", "concurrency_matrix.png", "各并发层的职责、隔离边界和一致性手段。", table_headers=["并发层", "实现", "共享范围", "保护机制"], table_rows=[
                ["角色进程", "simulator / trainee", "各自模型目录与运行态", "进程地址空间隔离"],
                ["潮流进程池", "spawn ProcessPoolExecutor", "同一 WEB 角色可共享 runner", "RLock、超时终止与重建"],
                ["后台线程", "clock/exchange/control", "模型级状态", "Event、RLock、run_lock、pending 集合"],
                ["HTTP 线程", "ThreadingHTTPServer", "同一服务内快照", "模型锁、定义 revision、原子替换"],
                ["WEB 客户端", "多个浏览器", "只读同一帧/独立 UI", "签名、revision、差量序号"],
            ], table_widths=[1.15, 1.7, 1.75, 1.9], source="simu/server.py::ManagedThreadingHTTPServer"),
            PageSpec("模拟台后台与潮流模拟进程", "每轮潮流采用一次性完整输入，跨进程求解完成后才在模型锁内发布结果。", "sim_flow_worker.png", "输入冻结、跨进程求解、结果返回和原子发布。", [
                "后台克隆内存 EBook，叠加天气、负荷、设备状态、控制指令、能力边界和 SOC，构造 SimulationConfig。",
                "runner.submit 后由子进程调用 simu_loop.run_once；返回 SimulationResult、runtime_stat_book、worker_pid 和耗时。",
                "默认超时 30 秒；超时或 BrokenProcessPool 会终止旧进程并重建执行器，避免卡死 worker 持续占用。",
                "只有求解成功且返回仍属于当前服务实例/模型代次时，后台才更新快照、历史量测和日志。",
            ], source="simu/power_flow_worker.py、simu/service.py::_compute_once_locked"),
            PageSpec("模拟台后台与模拟台 WEB 前端", "模拟台前端通过 REST/JSON 获取共享状态，支持多个客户端同时查看同一模型。", "sim_web_flow.png", "模拟台 WEB 读写路径和后台快照边界。", [
                "GET 读取配置、模型列表、快照、量测差量、历史、曲线、设备状态和日志；读取不推进时钟。",
                "POST 修改时钟、曲线、设置、故障和定义；后台校验 model_id 后在目标模型锁内提交。",
                "全量快照包含 clock、devices、measurements、definitions、settings 和 model_version；差量接口减少高频传输。",
                "客户端刷新失败只影响本地显示；后台不会因为单个浏览器断开而停止仿真或撤销其他客户端状态。",
            ], source="simu/server.py::_handle_api_get/_handle_api_post"),
            PageSpec("模拟台后台与学员台后台", "两后台通过交互链接建立显式数据契约：先同步定义，再接收运行帧，最后回送控制。", "teacher_trainee_flow.png", "模型发现、定义初始化、实时接收和控制回送。", [
                "交互链接返回 model_id、teacher_api_base、模型版本和 snapshot/command/delta/definition 路径。",
                "模型初始化下载定义包，校验签名后生成学员台本地定义；运行态不通过名称猜测重建角色。",
                "实时接收线程按周期拉取全量或差量，检查帧龄、同帧持续时间、run_id、定义 revision 和连接签名。",
                "学员台控制通过 /api/student/commands 回送；模拟台只接受合法来源、合法字段、未过期且当前代次的命令。",
            ], source="simu/server.py::_trainee_link_payload、simu/trainee_exchange.py"),
            PageSpec("学员台后台、WEB 前端与新能源控制单元", "学员台页面、实时交换层和控制工作单元共享模型级状态，但各自的耗时操作不持有大锁。", "trainee_control_flow.png", "学员台内部三层与模拟台命令入口的数据流。", [
                "WEB 前端读取 /api/trainee/snapshot、receive-state 和 renewable-control 状态；参数修改通过同一路径写入后台。",
                "实时交换层提供带 control_lease 的固定快照，包含 service_instance_id、receive_epoch、definition revision 和帧身份。",
                "控制工作单元先在后台线程收集输入，在线程池中计算；提交前用同一 generation guard 再次验证快照仍有效。",
                "停止接收、切换模型或服务退休会使租约失效；正在计算的结果即使完成也不能提交。",
            ], source="simu/trainee_exchange.py、simu/renewable_control.py::TraineeRenewableControlManager"),
            PageSpec("一致性、代次替换与 fail closed", "系统通过服务实例、接收代次、策略代次和定义修订建立端到端生命周期保护。", "generation_lifecycle.png", "strategy_id + generation 的生成、替换和撤销。", [
                "每轮自动策略携带固定 strategy_id 和递增 generation，并提交完整设备目标快照，而不是增量修补上一轮。",
                "模拟台收到新代次时先撤销同 strategy_id 的旧代次，再登记新代次；空目标快照也能明确撤销旧目标。",
                "停止、接收丢失、数据过期、模型切换、控制器退休和服务实例替换均撤销当前 generation。",
                "拓扑无效、量测缺失、设备边界无效、控制字段不明确或代次校验失败时，对相关设备/岛 fail closed。",
            ], callout="当前实现刻意不恢复设备名称猜测，也不依赖 dev_type 角色标记；点名中的设备类型只用于稳定命名空间。", source="simu/service.py::_cancel_strategy_generation、simu/renewable_control.py"),
        ]),
        ChapterSpec(3, "模拟台功能介绍", [
            PageSpec("主页：运行总览与能量流", "主页集中展示时钟、天气、交直流电源/负荷、储能、柴发、变流器和绿电指标。", "01_simulator_home.png", "模拟台主页及交直流微电网能量流。", [
                "顶部选择模型和仿真模式，使用启动、暂停、停止、单步、加速和减速控制仿真时钟。",
                "能量流卡片同时显示当前值、目标值、最大可发、SOC、台数和运行方向。",
                "DC/AC 变流器总值统一按 P_DC 显示；正值代表直流侧向交流侧送电。",
                "绿电指标按负荷减柴发计算，适合训练时快速判断新能源与储能共同承担的供电比例。",
            ], source="web/simulator 前端主页、simu/service.py 快照统计"),
            PageSpec("模型管理：新建、复制、导入和切换", "模型槽位将源定义与运行目录分离；切换模型时后台同时切换服务实例和生命周期。", "28_simulator_model_management.png", "模拟台模型管理窗口。", [
                "选择模型只改变当前浏览模型；正在运行的其他模型可由多模型时钟线程继续推进。",
                "新建/复制/导入会生成独立 model.e、control.e、meas.e、stat.e、curves 和 diagram。",
                "更新定义时使用 revision 做冲突检测，避免两个客户端以旧版本覆盖新修改。",
                "删除或替换模型会退休旧服务实例，关闭其潮流/控制相关状态，旧请求不得写回新实例。",
            ], source="simu/server.py::create_model_from_efile/import_definition_model、MultiModelSimulator"),
            PageSpec("电网模型：设备树、参数表与稳定身份", "电网模型页面按模型块展示设备和参数，不使用设备名称推断资源类型。", "02_simulator_grid_model.png", "设备树与分类型参数表。", [
                "左侧设备树用于过滤和定位；右侧表格按 ACNode、DCNode、Branch、Generator、Load、Storage、Converter 等模型块分页。",
                "设备稳定关系来自块类型、idx、端子节点和关联参数块中的显式索引；名称只用于显示和点名。",
                "参数编辑必须保留上下限、额定容量、控制模式和拓扑引用的一致性，否则运行时闭锁相关设备。",
            ], source="simu/server.py::_generated_model_artifacts、simu/model_semantics.py"),
            PageSpec("设备详情与 DC/AC 控制模式", "设备详情窗口用于核对端子、运行状态、控制方式、设值和真实功率边界。", "04_simulator_converter_detail.png", "DCACConverter 详情：ac_control_type、dc_control_type 与两端设值。", [
                "dc_control_type=P 时遥调 p_dc_set；dc_control_type=NONE 时遥调 p_ac_set；两侧均 NONE 时回退 p_dc_set。",
                "下发设值不得隐式修改 ac_control_type 或 dc_control_type，控制方式只由模型定义或显式模式操作改变。",
                "系统统计使用 P_DC，但详情可同时展示 p_ac/p_dc 物理端口值，二者符号相反并考虑损耗。",
                "若两台并联变流器控制模式不同，系统仍先在 P_DC 口径汇总，再逐台选择对应控制字段。",
            ], source="simu/device_roles.py、simu_loop.py::_set_value_target_column"),
            PageSpec("接线图：拓扑定位与设备状态", "接线图把模型节点、支路、开关、风光储柴荷和变流器映射到可视拓扑。", "05_simulator_diagram.png", "模拟台一次接线图。", [
                "颜色和开关状态反映当前运行态；点击设备可查看定义、量测和控制信息。",
                "接线图只是展示层，控制算法使用模型端子和稳定索引解析拓扑，不读取 SVG 文本或设备名称。",
                "死岛、断开路径和无参考源分量在后台拓扑解析后反馈到设备状态，相关自动控制应闭锁。",
            ], source="diagram.svg、simu/resource_topology.py、kernel topology"),
            PageSpec("曲线设置：风、光、温度与仿真尺度", "曲线页面按时/日/周/月/年仿真生成、编辑、扰动和保存环境边界。", "30_simulator_solar_curve.png", "太阳辐照曲线及逐点数据表。", [
                "日仿真通常按 1 分钟点读取；其他仿真尺度根据模式配置确定点数和曲线步长。",
                "风速进入风机最大可发曲线，辐照和温度进入光伏能力修正；页面值也作为气象遥测发布。",
                "扰动只改变选中曲线；保存后后台刷新曲线 revision，下一仿真步读取新值。",
            ], source="simu/service.py 曲线管理、simu_loop.py::wind_available_power/pv_available_power"),
            PageSpec("负荷曲线：逐设备需求与典型曲线", "交流负荷和直流负荷分别建曲线，仿真步按当前时刻读取并写入对应设备。", "31_simulator_load_curve.png", "直流负荷典型日曲线和可编辑数据。", [
                "负荷按设备逐项读取和插值，不把总负荷平均分配回设备；日志保留每个负荷分项。",
                "生成典型曲线可快速创建训练场景，扰动用于叠加波动，保存后成为下一轮计算输入。",
                "交流/直流负荷分别进入对应网络方程，跨侧平衡由网侧 DC/AC 变流器连接。",
            ], source="simu_loop.py::apply_load_model/apply_load_power_targets"),
            PageSpec("参数配置、故障与运行模式", "参数页面控制计算周期、噪声、初始 SOC 等运行参数；故障和模式用于场景编排。", "07_simulator_parameters.png", "模拟台运行参数配置。", [
                "计算周期决定后台墙钟多久触发一次检查；仿真时步和倍率共同决定每次推进的仿真分钟。",
                "设备故障可改变投退或开关状态，量测故障改变有效位/量测值；二者应分开使用和复盘。",
                "运行模式是成组状态/设值覆盖，仍通过显式设备身份匹配；不允许以名称关键字批量猜测类型。",
                "参数修改写入模型运行设置并在下次启动加载，多客户端应通过 revision 避免互相覆盖。",
            ], source="simu/web_runtime_settings.py、simu/service.py 故障/模式 API"),
            PageSpec("人工修改：定义修改、重试与恢复", "人工修改面向定义层或运行层的受控修改，页面保留修改记录和失败原因。", "08_simulator_manual_changes.png", "人工修改列表和处理状态。", [
                "定义修改请求带模型 revision；后台验证字段、比例参数和目标设备后生成新定义快照。",
                "失败修改可以重试或重置，但不会越过设备上下限、控制模式和拓扑结构校验。",
                "工作树/模型源文件与运行内存分离，运行叠加不应反向覆盖用户源模型中的控制方式。",
            ], source="simu/definition_editing.py、/api/definitions/manual-changes"),
            PageSpec("控制指令：当前有效、历史与取消", "控制指令页面用于查看人工、学员台和外部系统的遥控遥调记录及剩余有效期。", "09_simulator_commands.png", "模拟台控制指令列表。", [
                "指令按 remote_control 与 remote_adjustment 分组，记录来源、请求值、执行值、仿真时刻和过期时刻。",
                "自动指令按有效期或策略 generation 失效；人工 hold_until_cancelled 指令保持到显式取消。",
                "新策略 generation 完整替换上一代，不让旧设备目标因本轮未出现而继续残留。",
                "取消操作保留历史审计字段，不物理删除记录；下一潮流步不再叠加已取消指令。",
            ], source="simu/service.py::submit_student_commands、命令历史与取消逻辑"),
            PageSpec("实时量测：遥测、遥信与测点趋势", "量测页面使用同一测点定义展示 SCADA 当前值、有效位和历史变化。", "10_simulator_measurements.png", "实时量测表和测点跟踪曲线。", [
                "遥测包含功率、电压、电流、SOC 和天气；遥信包含投退、开合和状态量。",
                "测点顺序由 meas 定义固定，名称和值通过 definition_signature 校验；不依赖客户端逐点猜测。",
                "invalid/missing 点应先检查设备是否在死岛、定义是否引用不存在设备、测量类型是否适用。",
                "高频页面优先读取 measurements/delta，序号失配时回退全量快照。",
            ], source="simu/measurement_delta.py、simu/measurement_history.py"),
            PageSpec("运行日志：输入边界、控制响应和求解摘要", "运行日志把每个仿真步拆成环境/负荷、控制指令、控制响应和潮流计算等事件。", "11_simulator_logs.png", "运行日志列表及详细计算摘要。", [
                "环境/负荷日志给出曲线模式、目标分钟、风速、辐照、温度和负荷分项。",
                "控制响应日志给出有效指令数、设备投退/开关、设值、SOC 和新能源最大可发。",
                "潮流计算日志给出 iter、normF、量测更新/缺失、风光储柴荷统计和含损耗功率差额。",
                "筛选日志类型可快速定位计算不收敛、模型不一致、控制被拒绝或数据过期。",
            ], source="simu/service.py::_runtime_log、simu_loop.py::SimulationResult"),
            PageSpec("交互链接与标准训练流程", "交互链接将当前模型的发现、快照、命令和定义接口封装为可复用连接信息。", "29_simulator_interaction_link.png", "模拟台为当前模型生成的学员台交互链接。", [
                "教员选择模型并生成链接；学员台输入链接，先初始化模型，再启动接收。",
                "模拟台启动仿真后，学员台观察量测、接线和负荷；必要时先预览新能源策略。",
                "闭环控制启动后，模拟台控制指令页核对来源、generation、有效期和执行结果。",
                "训练结束先停止自动控制，再停止接收和模拟台时钟；确认当前自动 generation 已撤销。",
            ], callout="链接标记为“多学员台共享”时，多个学员台可读取同一模型；写控制仍在模拟台模型锁内按请求顺序登记。", source="/api/trainee-link、/api/student/commands"),
        ]),
        ChapterSpec(4, "模拟台潮流模拟技术", [
            PageSpec("一次仿真步的完整计算流程", "潮流不是独立按钮式计算，而是仿真步中天气、控制、拓扑、能力边界和量测发布的核心环节。", "simulation_cycle.png", "模拟台单步计算的四个阶段。", [
                "时钟线程判断是否到达计算时刻；单步操作直接触发一次相同流程。",
                "所有输入在提交潮流 worker 前冻结，避免浏览器写操作在求解中途改变边界。",
                "求解完成后先校验收敛，再更新 SOC、生成量测、历史和日志，最后一次性发布新快照。",
            ], source="simu/service.py::_advance/_compute_once_locked、simu_loop.py::run_once"),
            PageSpec("设备定义、关联参数与运行叠加", "潮流内核读取 EBook 模型块；运行时只在克隆模型上叠加状态和控制，不修改源定义。", "model_blocks.png", "潮流相关模型块、显式索引与运行叠加规则。", [
                "主设备块定义电气端子和状态，风机/光伏/储能参数块通过显式索引关联主发电机或变流器。",
                "支路、开关和变流器使用 i_node/j_node 或 AC/DC 端子索引建立稳定拓扑关系。",
                "weather、stat、control 的行按模型块与 idx/name 匹配，只覆盖模型已有且允许写入的列。",
                "不存在、重复或跨块冲突的身份不使用名称修复；运行时对相关设备闭锁并记录诊断。",
            ], source="EBook、simu_loop.py::apply_overlay_book、_embedded_device_define_book"),
            PageSpec("拓扑、母线合并和运行岛", "方程装配前先把结构模型转成当前运行拓扑，剔除断开、停运和无参考源的分量。", "topology.png", "结构图、运行图、母线和运行岛的构建过程。", [
                "零阻抗支路和闭合开关合并节点为计算母线；普通支路保留为导纳元件。",
                "AC 岛需要参考相角/平衡源，DC 岛需要电压参考；自动参考只在合格设备中选择。",
                "混联运行岛过滤同时考虑 AC、DC 与有效网侧变流器，避免死岛资源进入统计和控制。",
                "拓扑签名用于缓存数组映射；结构变化时重建，普通运行值变化可复用稀疏结构。",
            ], source="model/topology.py、model/ppc_topology.py"),
            PageSpec("交流支路方程", "交流线路、变压器和并联电纳统一装配到复数节点导纳矩阵 Ybus。", "ac_equations_math.png", "交流支路方程要素。", [
                "普通支路串联导纳 y=1/(r+jx)，两端各计入 b/2；变压器额外计入 tap 和 shift。",
                "从节点电压相量计算支路电流和端口复功率，端口损耗由两端功率之和得到。",
                "断开或死岛支路不进入有效 Ybus；零阻抗支路通过母线合并/附加约束处理，避免直接 1/0。",
            ], image_height=1.9, equations=[F["ac_admittance"], F["ac_node_power"]], source="lfcore/ac_lf.py::matpower_branch_stamp_vectorized、_prepare_ppc_y_matrix"),
            PageSpec("交流节点网络方程与控制节点", "交流潮流以节点有功/无功不平衡为残差，节点类型决定状态变量与方程集合。", "ac_equations_math.png", "交流节点网络方程要素。", table_headers=["节点类型", "已知量", "状态变量", "残差/控制"], table_rows=[
                ["PQ", "P、Q", "相角 θ、电压 V", "ΔP=0、ΔQ=0"],
                ["PV", "P、V", "相角 θ、无功 Q", "ΔP=0，V 固定；Q 受上下限校核"],
                ["Slack", "V、θ", "P、Q 由平衡决定", "提供岛内功率缺额和参考相角"],
                ["变流器电压控制端", "控制类型/设值", "由耦合变量决定", "附加电压或功率控制方程"],
            ], table_widths=[1.1, 1.7, 1.7, 2.0], image_height=1.45, equations=[F["ac_residual"]], source="lfcore/ac_lf.py::_fill_residual、_validate_acac_terminal_control"),
            PageSpec("直流支路与节点网络方程", "直流网络使用节点电压和电导矩阵，非参考节点满足有功平衡。", "dc_equations_math.png", "直流支路电流、端口功率、损耗和节点残差。", [
                "支路电流由电压差和电导决定；端口功率使用本端电压乘支路电流，损耗为 rI²。",
                "DCGenerator、DCLoad、储能和变流器端口功率共同进入节点指定注入。",
                "电压控制节点从状态变量中消去或用控制方程固定，其他节点电压由 Newton 迭代求得。",
                "DCDC 变流器可根据控制模式、两端电压和损耗关系补充功率/电压约束。",
            ], image_height=1.9, equations=[F["dc_branch"], F["dc_residual"]], source="lfcore/dc_lf.py::_eval_newton_terms、_dcdc_j_power_from_loss"),
            PageSpec("DC/AC 变流器方程与功率方向", "网侧变流器是 AC/DC 方程的耦合边界，系统统计与物理端口必须明确区分。", "converter_equations_math.png", "变流器端口功率、r1/r2 损耗方程和系统口径。", [
                "内核残差同时修改 AC 节点注入和 DC 节点注入，并增加功率耦合或电压控制方程。",
                "系统 P_DC 正值代表 DC→AC；在物理端口 p_dc>0、p_ac<0，损耗使 |送端|>|受端|。",
                "控制类型决定哪一端设值进入状态/残差；设值命令不能修改控制类型本身。",
                "并联变流器除各自上下限外，还可按额定容量设置共享比例约束。",
            ], image_height=2.0, equations=[F["converter"]], source="lfcore/hybrid_lf.py::_append_dcac_residuals、simu/device_roles.py"),
            PageSpec("混联全局状态、残差和网络方程", "HybridPowerFlowCalc 将 AC 子系统、DC 子系统和变流器耦合方程装配为一个全局非线性系统。", "newton_math.png", "混联状态向量、残差向量和 Newton 标准形式。", [
                "x=[θ_AC, V_AC, V_DC, x_converter]；具体维数取决于节点类型、运行岛和变流器控制模式。",
                "F=[F_AC(P,Q), F_DC(P), F_converter]，方程数与有效状态变量数一致。",
                "AC/DC 子求解器提供各自残差与雅可比；混联求解器追加耦合行列并使用统一收敛判据。",
            ], image_height=2.0, equations=[F["newton"]], source="lfcore/hybrid_lf.py::_split_x/get_f/get_jacobi"),
            PageSpec("Newton-Raphson、稀疏雅可比与线性求解", "程序每次迭代求解 J(x_k)Δx=F(x_k)，再按 x_(k+1)=x_k-Δx 更新，直到残差无穷范数满足容差。", "newton_math.png", "混联潮流 Newton-Raphson 标准公式。", [
                "首次 prepare 缓存节点映射、导纳、控制掩码和雅可比稀疏模式；后续迭代只刷新数值。",
                "全局雅可比由 AC、DC 和变流器块拼接，优先使用可用的稀疏直接求解器并保留 SciPy 回退。",
                "收敛日志中的 iter 与 normF 是首要质量指标；列零、矩阵奇异或超迭代上限都视为失败。",
                "求解失败不发布半成品结果；模拟台保留上一帧并记录错误，潮流子进程超时则重建。",
            ], image_height=2.0, equations=[F["newton"]], source="lfcore/hybrid_lf.py::_run_newton_raphson、_build_newton_system"),
            PageSpec("动态能力边界、SOC 与设备安全", "潮流计算前先将天气、SOC、额定值和控制指令投影到设备合法边界。", "capability_math.png", "风机、光伏和储能 SOC 的标准能力公式。", [
                "风机采用切入-额定-切出分段曲线；光伏按辐照、额定功率和温度系数计算最大可发。",
                "储能 p_set 正值放电、负值充电；SOC 达下限禁止放电、达上限禁止充电，并按效率积分更新。",
                "柴发、风光、储能和变流器均先校验真实 p_min/p_max 或充放电/传输上限，再应用设值。",
                "拓扑无效或边界缺失时不使用默认大上限代替安全数据，相关设备 fail closed。",
            ], image_height=1.7, equations=[F["wind_pv"], F["soc"]], source="simu_loop.py::apply_device_capability_limits_book/update_storage_soc_book"),
            PageSpec("量测生成、跨进程隔离与日志校核", "潮流结果经过测点映射、噪声和故障处理后发布为可供 WEB/学员台/外部系统读取的统一帧。", "measurement_math.png", "真值映射、SCADA 噪声模型和有效位公式。", [
                "real 保存潮流真值，scada 叠加噪声与量测故障；两者都保持 meas 定义顺序和有效位。",
                "存储 SOC、运行状态和天气量可以作为非电气测点并入同一名称/值序列。",
                "跨进程返回携带 worker PID、计算耗时和往返耗时；健康接口展示超时、重启和最后原因。",
                "日志功率差额包含网络与变流损耗，不要求严格为 0；异常偏大时核对边界、符号和缺失测点。",
            ], image_height=2.0, equations=[F["measurement"]], source="simu_loop.py::build_real_rows_from_data/add_noise_to_rows、simu/power_flow_worker.py"),
        ]),
        ChapterSpec(5, "模拟台对外 WEB 接口", [
            PageSpec("接口发现与访问入口", "外部程序不应硬编码全部路径，而应先读取交互链接返回的 external_api。", "external_api_map.png", "模拟台外部 WEB 接口分组。", code='''GET http://127.0.0.1:8710/api/trainee-link?model_id=%E7%A7%A6%E5%B2%AD%E7%AB%99

返回关键字段：
model_id, teacher_api_base, model_version,
 external_api.devices, realtime_inputs, telemetry_names, telemetry_values,
selected_telemetry_values, measurement_history,
control_names, control_execute''', bullets=[
                "model_id 必须 URL 编码；后续路径优先使用链接返回值，避免服务升级时路径漂移。",
                "GET 接口用于发现和读取，POST 接口用于选择查询、历史查询和控制执行。",
                "所有外部接口均返回 JSON；错误也使用统一 JSON 结构并带 HTTP 状态码。",
            ], source="docs/external-web-api.md、simu/server.py::_trainee_link_payload"),
            PageSpec("模型版本与点表一致性", "外部系统必须把模型定义版本和点表顺序当作帧接收前置条件。", "version_handshake.png", "交互链接、名称和数值帧的签名握手。", [
                "model_version.signature 不一致：丢弃当前帧，重新读取交互链接和全部定义。",
                "definition_signature 不一致：丢弃值帧，重新读取遥测/遥信或控制名称列表。",
                "名称和值长度不一致：告警并丢弃；全量值数组严格按名称数组顺序排列。",
                "运行值、时钟和控制变化不改变 model_version；模型/量测/控制定义变化才改变 revision 和签名。",
            ], source="/api/external/* 响应、simu/measurement_delta.py"),
            PageSpec("全部设备与拓扑接口", "设备接口提供重建网络拓扑和设备详情所需的静态、运行和当前值信息。", "02_simulator_grid_model.png", "设备定义在模拟台电网模型页面中的展示。", code='''GET /api/external/devices?model_id=秦岭站

{
  "devices": [{
    "device_type": "DCACConverter",
    "name": "DCAC变流器-1",
    "topology": {...}, "parameters": {...},
    "parameter_blocks": [...], "state": {...},
    "values": [...], "control_values": [...]
  }],
  "topology": {"nodes": [...], "connections": [...]}
}''', bullets=[
                "topology.nodes/connections 可直接用于外部拓扑图，不要通过名称文本推断连接。",
                "parameters 与 parameter_blocks 分开，便于读取主设备和风光储关联参数。",
                "state 包含投退、开合、死岛、控制模式和 SOC；values/control_values 提供当前点值。",
            ], source="/api/external/devices"),
            PageSpec("实时环境与负荷输入：发现、字段与单点写入", "外部采集程序或上位系统可把天气和交直流负荷写入模拟台运行曲线，供后续潮流时步读取。", "realtime_inputs_flow.png", "从接口发现、契约读取、原子写入到潮流生效的完整链路。", code='''GET /api/external/realtime-inputs?model_id=秦岭站

POST /api/external/realtime-inputs?model_id=秦岭站
Content-Type: application/json
{
  "start_time": "00:10:00",
  "point_count": 1,
  "point_interval_seconds": 60,
  "weather": {"wind_speed_mps": 9.6,
              "solar_irradiance_w_m2": 650},
  "loads": {"ACLoad:交流负荷-1": 120,
            "DCLoad:直流负荷-1": 40}
}''', bullets=[
                "先从 trainee-link.external_api.realtime_inputs 取得路径；GET 返回天气字段、单位、标准负荷键、下一目标点和曲线采样契约。",
                "天气字段包括风速、太阳辐射、气温、气压和湿度；负荷键使用 ACLoad:设备名 或 DCLoad:设备名。",
                "兼容请求可省略时标并更新下一未求解点；新接入程序应显式提供 start_time、point_count 和数据点间隔。",
                "接口只在模拟台提供；学员台对 GET/POST 均返回 404，不代理写入模拟台。",
            ], image_height=1.9, source="docs/external-web-api.md §3.1、service.py::external_realtime_input_schema/apply_external_realtime_inputs"),
            PageSpec("实时环境与负荷输入：批量时序、原子校验与响应", "批量接口支持逐点对象 points 和字段序列 series，两种格式共用严格的时标、间隔和整帧原子性规则。", "realtime_inputs_contract.png", "批量时序输入必须先通过完整契约校验，再一次写入运行曲线。", code='''POST /api/external/realtime-inputs?model_id=秦岭站
{
  "start_time": "00:00:00",
  "point_count": 3,
  "point_interval_seconds": 60,
  "series": {
    "weather": {
      "wind_speed_mps": [8.0, 8.2, 8.5],
      "solar_irradiance_w_m2": [0, 20, 50]
    },
    "loads": {
      "ACLoad:交流负荷-1": [100, 105, 110]
    }
  }
}''', table_headers=["处理环节", "接口规则与结果"], table_rows=[
                ["时序定位", "start_time/绝对时标是权威依据；间隔必须为当前曲线采样间隔的正整数倍"],
                ["数组校验", "points 长度和每点字段集合一致；series 每条数组长度等于 point_count"],
                ["数值与设备", "风速、辐照和负荷非负，湿度 0..100，气压>0；负荷设备必须存在且键唯一"],
                ["原子性", "任一字段非法返回 HTTP 400，整帧不写入，curves 和 curve_revision 均保持不变"],
                ["成功响应", "返回 update_mode、updated_indices、起止时刻、curve_revision、curve_boundary 和 applies_on_next_power_flow"],
                ["并发语义", "潮流正在计算时等待该轮结束再写入；请求不推进仿真时钟，也不修改源 curves.json"],
            ], table_widths=[1.25, 5.25], image_height=1.45, source="docs/external-web-api.md §3.1、tests/test_external_realtime_inputs.py"),
            PageSpec("遥测/遥信名称与全量数值", "名称和值分开获取，适合高频读取时缓存点表并只请求紧凑值数组。", "measurement_pipeline.png", "测点定义顺序贯穿真值、SCADA 和外部接口。", code='''GET /api/external/telemetry/names?model_id=秦岭站
GET /api/external/telemetry/values?model_id=秦岭站

names: telemetry_names[], signal_names[], definition_signature
values: telemetry_values[], signal_values[],
        telemetry_valid[], signal_valid[], simu_time

兼容别名：yc_names/yx_names/yc_values/yx_values''', bullets=[
                "自动点名采用 设备类型.设备名称.点类型，设备类型在这里是命名空间，不是控制角色推断依据。",
                "valid=0 表示本帧点值不可用；调用方不得把 null/invalid 当 0 参与控制。",
                "高频程序应缓存名称和索引，以数组下标读取值；签名变化时重建缓存。",
            ], source="/api/external/telemetry/names、/values"),
            PageSpec("指定遥测/遥信查询", "只关心少量点时使用选择查询，返回顺序与请求顺序一致，并显式报告缺失点。", "external_api_map.png", "选择查询属于紧凑读取接口。", code='''POST /api/external/telemetry/values/query?model_id=秦岭站
Content-Type: application/json

{
  "telemetry_names": [
    "Environment.weather.WIND_SPEED",
    "ACLoad.交流负荷-1.P_LOAD"
  ],
  "signal_names": [
    "ACGenerator.柴油发电机-1.run_stat"
  ]
}''', bullets=[
                "返回 telemetry_found、signal_found 和 missing；不存在的点保持原位置，值为 null、valid=0。",
                "调用方可以稳定地把返回数组映射回请求数组，不需要二次名称排序。",
                "点类型不匹配也按 missing 处理，避免把遥信点误当遥测点。",
            ], source="/api/external/telemetry/values/query"),
            PageSpec("历史遥测与遥信曲线", "历史接口支持累计时间、仿真时刻和采样间隔，返回时间优先二维数组。", "external_api_map.png", "历史查询与实时查询共用点表签名。", code='''POST /api/external/telemetry/history/query?model_id=秦岭站
{
  "start_time": "00:00:00",
  "end_time": "02:00:00",
  "interval_seconds": 300,
  "telemetry_names": ["Environment.weather.WIND_SPEED"],
  "signal_names": ["ACBreak.盒型开关-1.status"]
}

telemetry_values[time_index][point_index]''', bullets=[
                "累计分钟/秒参数优先于 HH:MM:SS，适合周、月、年仿真跨天查询。",
                "请求间隔小于原始采样间隔时使用目标时刻之前最近一帧，不对遥测线性造值。",
                "单次最多 256 个名称、10000 个采样点；历史只覆盖当前 run_id，重新开始后清空。",
            ], source="/api/external/telemetry/history/query、simu/measurement_history.py"),
            PageSpec("遥调/遥控名称", "控制名称接口返回当前模型允许外部下发的遥调和遥控点。", "09_simulator_commands.png", "模拟台控制指令页面可核对外部控制的登记结果。", code='''GET /api/external/controls/names?model_id=秦岭站

{
  "remote_adjustment_names": [
    "ACGenerator.柴油发电机-1.p_set"
  ],
  "remote_control_names": [
    "ACBreak.盒型开关-1.status"
  ],
  "definition_signature": "<sha256>"
}

兼容别名：yt_names / yk_names''', bullets=[
                "只允许名称表中的点下发；设备存在但字段不适用时仍视为未找到。",
                "变流器遥调字段按 ac_control_type/dc_control_type 选择，不按 p_ac_set/p_dc_set 是否存在来抢优先级。",
                "控制定义签名变化时必须重新获取名称，避免旧字段在模型切换后误写。",
            ], source="/api/external/controls/names、simu/device_roles.py"),
            PageSpec("提交遥调、遥控与代次控制", "执行接口支持普通外部 EMS 控制，也支持人工保持、自动有效期和策略代次元数据。", "strategy_snapshot.png", "控制快照登记、替换、执行和撤销。", code='''POST /api/external/controls/execute?model_id=秦岭站
{
  "remote_adjustment_names": ["ACGenerator.柴油发电机-1.p_set"],
  "remote_adjustment_values": [60.0],
  "remote_control_names": ["ACBreak.盒型开关-1.status"],
  "remote_control_values": [0],
  "valid_for_minutes": 10,
  "source": "external-ems"
}''', bullets=[
                "名称和值数组必须等长；每项返回 found、accepted、active、reason、执行值和失效时刻。",
                "默认按自动指令处理并有有效期；人工长期保持应传 command_origin=manual、hold_until_cancelled=true。",
                "自动策略可带 strategy_id、generation、replace_generation，以完整代次替换旧策略。",
            ], source="/api/external/controls/execute、simu/service.py 命令登记"),
            PageSpec("PowerShell 完整调用示例", "下面示例完成发现、签名校验、历史读取和遥调下发。", "version_handshake.png", "外部程序应先校验版本，再消费值或下发控制。", code='''$link = Invoke-RestMethod `
  "http://127.0.0.1:8710/api/trainee-link?model_id=秦岭站"
$base = $link.teacher_api_base
$names = Invoke-RestMethod ($base + $link.external_api.telemetry_names)
$frame = Invoke-RestMethod ($base + $link.external_api.telemetry_values)
if ($names.model_version.signature -ne $frame.model_version.signature -or
    $names.definition_signature -ne $frame.definition_signature) {
  throw "模型或点表版本不一致，本帧已丢弃"
}
$body = @{ remote_adjustment_names=@("ACGenerator.柴油发电机-1.p_set");
  remote_adjustment_values=@(60.0); valid_for_minutes=10;
  source="external-ems" } | ConvertTo-Json
Invoke-RestMethod -Method POST -ContentType "application/json" `
  -Uri ($base + $link.external_api.control_execute) -Body $body''', source="docs/external-web-api.md"),
            PageSpec("Python/cURL、错误处理与网络安全", "生产接入应设置超时、处理 HTTP/JSON 错误、验证 accepted，并限制网络暴露面。", "external_api_map.png", "外部接口发现、读取、历史和控制的统一入口。", code='''import requests
base = "http://127.0.0.1:8710"
link = requests.get(base + "/api/trainee-link",
                    params={"model_id": "秦岭站"}, timeout=5).json()
frame = requests.get(link["teacher_api_base"] +
                     link["external_api"]["telemetry_values"],
                     timeout=5).json()

curl -X POST "http://127.0.0.1:8710/api/external/telemetry/values/query?model_id=秦岭站" \\
  -H "Content-Type: application/json" \\
  -d '{"telemetry_names":["Environment.weather.WIND_SPEED"]}' ''', bullets=[
                "400 表示请求字段/值错误，404 表示模型或路由不存在，409 常用于定义 revision 冲突，500 表示后台异常。",
                "任何 accepted=false 都必须读取 reason；网络 200 不等于每个控制点都已被接受。",
                "服务无内置认证，局域网部署应使用反向代理认证、TLS、防火墙来源限制和审计日志。",
                "多客户端可并发 GET；控制 POST 在模型锁内登记，但外部系统仍应使用策略代次避免相互覆盖。",
            ], source="simu/server.py::JsonApiError/_cors"),
        ]),
        ChapterSpec(6, "学员台功能介绍", [
            PageSpec("主页与接收状态", "学员台主页展示模拟台数据源、模型、量测可用率、能量流、环境边界和当前有效指令。", "12_trainee_home.png", "学员台未接收时的主页和数据源信息。", [
                "先确认数据源链接和模型 ID，再执行模型初始化；定义一致后才能启动实时接收。",
                "接收状态区区分已冻结、接收中、通讯失败、数据过期等状态，并显示可用量测点数。",
                "主页能量流与模拟台保持同一统计口径，DC/AC 总功率/目标使用 P_DC。",
                "多个浏览器客户端可查看同一学员台模型；接收和控制状态由后台共享。",
            ], source="web/trainee 首页、/api/trainee/receive-state"),
            PageSpec("模型管理与模型初始化", "学员台本地模型是模拟台定义的受控副本，运行帧与本地定义通过版本签名绑定。", "27_trainee_model_management.png", "学员台模型管理窗口。", [
                "新建模型槽位后输入交互链接，初始化流程下载定义包并验证模型、测点和控制定义。",
                "初始化不会根据设备名称补角色；显式模型块、稳定索引和拓扑关系必须有效。",
                "重新初始化或切换模型会停止当前接收，撤销自动策略 generation，并退休旧控制状态。",
            ], source="/api/trainee/model-initialize、simu/server.py::_handle_trainee_model_initialize"),
            PageSpec("电网模型与接线图", "学员台使用本地定义显示设备参数，用接收的运行帧叠加状态和量测。", "15_trainee_diagram.png", "学员台接线图。", [
                "电网模型表按模型块分页；运行值来自接收快照，不把模拟台源模型文件直接映射为可写文件。",
                "接线图用于理解当前开关、死岛和能量路径；自动策略仍以后台拓扑解析结果为准。",
                "若定义签名与接收帧不一致，学员台停止接收并要求重新初始化，防止错位点表参与控制。",
            ], source="13_trainee_grid_model.png、simu/trainee_exchange.py::_merge_runtime_snapshot_with_local_definitions"),
            PageSpec("曲线显示与实时量测", "学员台曲线用于观察收到的环境、功率、SOC 和控制趋势，不直接修改模拟台曲线。", "17_trainee_measurements.png", "学员台实时量测表。", [
                "实时量测继承模拟台测点顺序和有效位；接收层可用差量帧更新本地快照。",
                "历史曲线按当前 run_id 分段，模型切换、重新开始或生命周期变化后不拼接旧段。",
                "控制页面的趋势数据来自固定控制快照和计划结果，便于对比实时值与目标值。",
            ], source="simu/trainee_exchange.py::measurement_delta/measurement_history"),
            PageSpec("控制指令与人工修改", "学员台既可查看/发起普通遥控遥调，也可由新能源控制生成自动策略。", "18_trainee_commands.png", "学员台控制指令页面。", [
                "人工遥控/遥调应先核对设备、点名、当前值、边界和控制模式，再提交到模拟台命令入口。",
                "人工修改页面用于本地定义/状态相关操作，接收中对关键定义修改会被限制，避免运行帧错配。",
                "自动新能源策略使用独立来源和 generation，不覆盖人工指令的审计属性；冲突按模拟台有效指令规则处理。",
            ], source="/api/trainee/commands、/api/student/commands"),
            PageSpec("参数配置、运行日志与通讯故障", "参数页控制接收刷新、超时和本地显示；日志页记录连接、接收、命令和控制生命周期。", "21_trainee_logs.png", "学员台运行日志。", [
                "接收刷新周期、请求超时、帧龄上限和同帧上限决定何时判定通讯或数据过期。",
                "通讯失败连续达到阈值后停止接收，自动控制前置条件失效并撤销当前策略代次。",
                "日志应结合模拟台状态检查：模拟台未启动、模型不一致、链接失效和定义签名变化是常见原因。",
            ], source="simu/trainee_exchange.py::_runtime_settings_for_service/receive_status"),
            PageSpec("新能源控制主页面", "新能源控制页面把运行参数、侧别统计、设备策略、曲线和日志组合在一个工作区。", "22_trainee_renewable_control.png", "新能源控制主页面。", [
                "左侧选择开环/闭环、控制周期和参数；接收前置条件不满足时启动和单次下发按钮禁用。",
                "统计指标分交流侧、直流侧和系统总；系统变流器指标统一按 P_DC。",
                "设备分类包含交流/直流风光、跟网/构网储能、柴发和 ACDC 变流器，并显示可用边界、目标和拓扑状态。",
                "曲线用于观察功率与 SOC，日志用于审计每轮决策、告警、generation 和下发结果。",
            ], source="web/trainee renewable control、/api/trainee/renewable-control"),
            PageSpec("系统总计、变流器目标与控制参数", "侧别和系统统计与逐设备目标使用同一计划结果，避免不同控制模式下重复或漏算。", "24_renewable_acdc_targets.png", "ACDC 设备分类及 P_DC 总目标口径。", [
                "两台变流器可一台控 DC、一台控 AC，也可都控同侧；统计先统一到 P_DC，再逐台换算并选择设值字段。",
                "并联功率按容量和剩余裕度合理分配，设备上下限或控制字段无效时仅闭锁对应设备/组。",
                "控制参数包含周期、指令有效期、风光/储能步长、构网储能/柴发保护带、SOC 死区、优化权重和容差。",
            ], source="23_renewable_system_totals.png、26_renewable_parameters_dialog.png"),
            PageSpec("学员台标准训练流程与故障恢复", "训练流程应把定义一致性、接收新鲜度、策略预览、闭环下发和代次撤销连成闭环。", "14_trainee_receiving_home.png", "通讯异常窗口可用于核对模拟台状态、链接和定义。", [
                "1. 选择本地模型并初始化；2. 启动模拟台；3. 启动接收并确认帧龄、量测可用率和模型一致。",
                "4. 在开环模式预览策略，检查弃电、柴发、SOC、保护带、步长和变流器方向；5. 再切闭环。",
                "6. 观察模拟台控制指令与下一潮流帧响应；7. 停止控制时确认当前 generation 已撤销。",
                "通讯失败先检查模拟台是否运行，再检查交互链接、模型 ID、定义签名、防火墙和请求超时。",
            ], source="学员台接收与新能源控制生命周期"),
        ]),
        ChapterSpec(7, "学员台新能源优先控制技术", [
            PageSpec("控制目标、优先级与输入快照", "算法在设备安全和拓扑有效前提下，以最小新能源弃电、最低柴发出力为主要目标。", "renewable_inputs.png", "固定控制快照、显式资源、运行拓扑和动态边界。", [
                "主要目标：最小化 Σ(P_avail-P_renewable)；在可比解中继续压低 ΣP_diesel。",
                "安全先验：设备功率/SOC 上下限、强制充放、动态降额、构网储能/柴发保护带和调节步长。",
                "平衡与拓扑：每个 AC/DC 运行岛和 DC 传输组独立建模，通过有效网侧变流器耦合。",
                "输入快照绑定 model revision、service_instance_id、receive_epoch、run_id 和帧身份，计算中不读取漂移数据。",
            ], source="simu/renewable_control.py::calculate_renewable_control_plan"),
            PageSpec("显式资源、拓扑解析与 fail closed", "资源类型来自模型块和显式关联，不使用名称关键字，也不依赖 dev_type 角色标记。", "topology.png", "资源拓扑从结构图叠加运行状态后形成可控集合。", [
                "风光储通过关联参数块和主设备稳定索引解析；柴发、网侧变流器按明确模型块、控制模式和端子识别。",
                "运行拓扑计算 AC/DC 分量、资源接入母线、路径、DC 传输组和网侧变流器集合。",
                "缺失端子、重复稳定身份、死岛、无有效量测、边界无效或控制模式不明确时闭锁相关设备/岛。",
                "失效范围尽量局部化：一个岛的问题不应让无关健康岛停止优化。",
            ], source="simu/resource_topology.py、simu/renewable_control.py::_apply_grid_forming_fail_closed_scopes"),
            PageSpec("三阶段计算总览", "优化按字典序分三阶段，确保经济目标不会破坏安全可行性或最小功率失衡。", "three_stage.png", "阶段 1 安全可行域、阶段 2 最小失衡、阶段 3 新能源优先经济目标。", [
                "每个拓扑岛独立建立变量和等式矩阵 A，变流器同时在 AC/DC 平衡行中出现。",
                "若阶段 1 已找到精确平衡解，阶段 2 的最小松弛为 0；否则显式计算 δ* 并保留告警。",
                "阶段 3 固定 δ* 后优化弃电、柴发和调节平滑，不能用经济收益换取更大失衡。",
            ], source="simu/renewable_optimization.py::_solve_island"),
            PageSpec("阶段 1：可行域与安全约束", "第一阶段的任务不是追求经济性，而是把所有设备约束投影成可信的优化边界并判断精确平衡是否可行。", "stage1_math.png", "阶段 1 的标准可行性优化模型。", table_headers=["要素", "定义"], table_rows=[
                ["输入", "固定快照、拓扑岛、当前功率、可发能力、SOC、额定值、控制模式、步长和保护参数"],
                ["变量", "风光/柴发/储能目标 P_i；网侧变流器统一 P_DC 目标"],
                ["目标", "标准线性可行性问题 min 0；仅检查 A·P=A·P_current 与并联等式在活动安全边界内是否可行"],
                ["边界", "hard lower/upper 与 normal step 交集；SOC 强制方向、降额和保护带先行"],
                ["输出", "active_lower/upper、可行初值；必要时记录 renewable safety step override"],
            ], table_widths=[1.25, 5.25], image_height=1.25, equations=[F["safety_bounds"], F["stage1"]], source="renewable_optimization.py::_renewable_variable/_diesel_variable/_storage_variable/_solve_island"),
            PageSpec("阶段 2：最小功率平衡松弛", "当精确平衡在当前安全边界内不可行时，第二阶段计算不可避免的最小 AC/DC 功率失衡。", "stage2_math.png", "阶段 2 的标准松弛定义、一般二次目标和 DC/AC 字典序线性规划。", bullets=[
                "DC 优先只在精确 AC/DC 平衡已证明不可行且存在网侧变流器时启用，不改变正常经济目标。",
                "输出 δ* 和对应可行 P*；数值不收敛但边界有效时保留可行回退并标注状态。",
            ], image_height=2.0, equations=[F["stage2"]], source="renewable_optimization.py::dc_priority_minimum_delta/solve_with_bounds"),
            PageSpec("阶段 3：新能源优先经济目标", "第三阶段在不恶化 δ* 的等式面上选择最小弃电、最低柴发且调节平滑的设备目标。", "stage3_math.png", "阶段 3 的标准目标函数、固定最小失衡约束、并联等式与安全边界。", bullets=[
                "w_r、w_d 为主要线性目标权重；平方权重用于平滑和等价解选择，不应压倒主要优先级。",
                "求解后在保持等式和边界的条件下按 SOC 重新分配并联储能，降低过充/过放风险。",
                "输出岛级目标、弃电、柴发、δ、迭代次数和状态；失败岛不生成可执行策略。",
            ], image_height=2.0, equations=[F["stage3"]], source="renewable_optimization.py::base_objective、_rebalance_storage_targets_by_soc"),
            PageSpec("SOC、强制动作、降额、步长与保护带", "安全约束在三阶段之前构造，并贯穿所有阶段。", "soc_derating.png", "储能 SOC 边界与充放电降额。", [
                "SOC≤下限：禁止继续放电；进入低 SOC 强制区时要求向充电方向修正。",
                "SOC≥上限：禁止继续充电；进入高 SOC 强制区时要求向放电方向修正。",
                "充/放功率先受额定功率限制，再乘 SOC 降额曲线因子，并与本轮最大调节步长求交。",
                "构网储能和柴发保护带用于防止计划贴近极限；保护修正优先于普通步长，必要覆盖会被显式记录。",
            ], image_height=1.45, equations=[F["safety_bounds"]], source="renewable_control.py::_storage_rows/_grid_storage_target_margins、renewable_optimization.py::_storage_variable"),
            PageSpec("DC/AC 变流器与并联合理分配", "变流器既是 AC/DC 平衡变量，也是控制字段和统计口径最容易出错的设备。", "parallel_converters.png", "并联容量分配、P_DC 统计、端口换算和字段选择。", [
                "系统总功率/总目标统一 P_DC，避免两台控制模式不同导致一台按 p_ac、一台按 p_dc 相加。",
                "同 AC/DC 分量间的并联变流器按额定可分配容量保持比例；设备饱和后按剩余裕度再分配。",
                "dc_control_type=P 生成 p_dc_set；dc_control_type=NONE 生成 p_ac_set；双 NONE 回退 p_dc_set。",
                "直流送交流显示正值，但若命令字段为 p_ac_set，实际写入的是负数端口设值；界面仍显示换算后的正 P_DC。",
            ], image_height=1.35, equations=[F["converter_dispatch"]], source="renewable_optimization.py::_converter_variable/parallel_matrix、simu/device_roles.py"),
            PageSpec("策略输出、完整快照与 generation", "优化结果经过统一校验后才转换为可下发的设备命令快照。", "strategy_snapshot.png", "计划结果到模拟台有效控制指令的生命周期。", [
                "输出包含 command_rows、侧别/系统指标、岛级优化状态、告警、决策细节和趋势点。",
                "每轮命令携带 strategy_id=renewable_priority、generation 和 replace_generation=true，并包含本轮全部目标。",
                "模拟台撤销旧 generation 后登记新目标；旧代次不会在新代次缺少某设备时继续残留。",
                "停止、接收丢失、数据过期、模型切换和控制器退休时，学员台发送当前 generation 撤销。",
            ], source="renewable_control.py::_command_payload、service.py::_cancel_strategy_generation"),
            PageSpec("开环、闭环、数据新鲜度与合理性评估", "同一算法既支持开环预览，也支持闭环下发；差别只在是否允许提交，不在约束和目标。", "25_renewable_strategy_log.png", "新能源控制日志页用于审计计算、告警和下发结果。", [
                "开环：固定快照计算并展示目标，不提交命令；闭环：提交前再次验证 generation、接收状态和帧新鲜度。",
                "同一仿真帧不重复下发同 generation；正在发送或已有后台周期时，新周期返回 busy/cancelled 状态。",
                "合理性检查包括功率/SOC 上下限、强制充放、弃电最小、柴发最低、保护带、步长、降额和并联分配。",
                "任何 invalid/failed/balance slack/step override 都进入日志和 metrics；不能把有告警的可行回退误报为严格最优。",
            ], callout="评价顺序：先检查安全约束和拓扑有效，再检查功率平衡松弛，最后比较弃电、柴发和调节平滑目标。", source="renewable_control.py::run_once/state、renewable_optimization.py::IslandOptimizationResult"),
        ]),
        ChapterSpec(8, "总结与展望", [
            PageSpec("能力总结与验收要点", "本平台形成了从场景构造、混联潮流、量测发布到学员控制和外部集成的闭环。", "outlook.png", "当前手册覆盖的操作、算法和接口能力。", table_headers=["验收域", "通过标准"], table_rows=[
                ["模拟台", "多模型、曲线/故障/模式、时钟、潮流、量测、日志和交互链接可用"],
                ["并发", "模拟台/学员台独立进程，多 WEB 客户端并发，潮流 worker 超时可恢复"],
                ["潮流", "设备/拓扑/AC/DC/变流器方程明确，收敛、限值、SOC 和量测链路可追溯"],
                ["接口", "版本握手、设备、全量/选择/历史量测、控制名称和执行示例可复现"],
                ["学员台", "初始化、接收、显示、控制、日志和故障恢复流程完整"],
                ["新能源控制", "三阶段输入输出明确，安全、弃电、柴发、步长、保护带、降额、并联均可审计"],
            ], table_widths=[1.45, 5.05], source="本手册各章与当前代码实现"),
            PageSpec("展望：进程化、安全与工程化演进", "现有架构已经具备清晰的数据契约和生命周期边界，可在不改变前端语义的前提下继续增强。", "outlook.png", "后续演进方向。", [
                "将新能源控制工作单元从学员台进程内线程池进一步拆成独立 OS 进程，沿用固定快照、generation guard 和命令快照契约。",
                "为外部 WEB 接口增加认证授权、TLS、来源配额、幂等键和更细粒度审计，避免开放网络中的未授权控制。",
                "扩展潮流工作池的按模型排队、公平调度、资源监控和热备 worker，提高多模型并发吞吐。",
                "持续用随机工况审计三阶段优化：记录严格最优、可行回退、平衡松弛、步长覆盖和 fail-closed 覆盖率。",
                "把模型 schema、控制字段、接口 JSON 和策略 metrics 版本化，形成可自动验证的兼容性契约。",
            ], callout="最终原则不变：模型身份显式、拓扑关系稳定、设备边界真实、数据生命周期可验证、失效行为 fail closed。", source="当前架构与建议演进方向"),
        ]),
    ]


def build_document() -> Path:
    generate_diagrams()
    doc = Document()
    setup_styles(doc)
    front_matter(doc)
    chapters = build_chapters()
    for chapter in chapters:
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        configure_section(section, f"第 {chapter.number} 章  {chapter.title}")
        for index, page in enumerate(chapter.pages):
            add_manual_page(doc, chapter, page, first_in_section=index == 0)
    core = doc.core_properties
    core.title = "极地微电网模拟台与学员台用户使用手册"
    core.subject = "模拟台、学员台、混联潮流、外部WEB接口与新能源优先控制"
    core.author = "极地微电网仿真与控制平台项目组"
    core.keywords = "模拟台, 学员台, 潮流, 新能源优先控制, WEB接口, AC/DC"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build_document()
    print(json.dumps({"output": str(output), "size": output.stat().st_size}, ensure_ascii=False))
